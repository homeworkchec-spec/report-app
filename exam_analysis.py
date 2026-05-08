"""
시험지 분석 자동화 모듈 — 블로그 발행용

핵심 흐름:
    §1) 이미지 업로드 → GPT-4o Vision OCR
        - 메타 + 문항 메타데이터 + 지문발췌 + 선지
        - 각 문항의 페이지 인덱스 + 세로 영역(top_ratio/bottom_ratio)
        - 원본 이미지를 세션에 보관 (크롭 재사용)
    §2) 메타 확인/수정 (인라인)
    §3) [블로그용 분석 보고서 생성] 단일 버튼
        - 어려운 문항 깊이 분석 (함정 / 풀이 / 대비)
        - 어려운 문항만 자동 크롭(이미지) — 문제 + 정답 근거 영역만
        - 기·승·결 본문 + 차트 캡션 (단일 LLM 호출, JSON)
        - 차트 4종 → 블록 리스트로 변환
    §4) 블록 편집기
        - 텍스트 수정, 이미지 추가/삭제/교체, 순서 변경
        - 어려운 문항 블록은 크롭 영역 슬라이더로 미세조정
    §5) 변경사항 적용 → PNG + Word + Text 재생성 → 다운로드

테마(차트 색감)는 사이드바에서 mono / editorial / vivid 선택.
"""

from __future__ import annotations

import base64
import io
import json
import platform
import random
import re
import uuid
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from typing import Any

import matplotlib
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
import numpy as np
import pandas as pd
import streamlit as st
from PIL import Image, ImageDraw, ImageFont

import openai

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from design_tokens import chart_palette, ThemeName


# ─────────────────────────────────────────────────────────────
# Korean font setup
# ─────────────────────────────────────────────────────────────
def _setup_korean_font() -> tuple[str | None, str | None]:
    candidates: list[Path] = []
    repo_font = Path(__file__).parent / ".fonts"
    if repo_font.exists():
        candidates += list(repo_font.glob("*.ttf")) + list(repo_font.glob("*.otf"))
    if platform.system() == "Windows":
        for n in ("malgun.ttf", "malgunbd.ttf", "NanumGothic.ttf"):
            p = Path("C:/Windows/Fonts") / n
            if p.exists():
                candidates.append(p)
    elif platform.system() == "Darwin":
        for n in ("AppleSDGothicNeo.ttc", "AppleGothic.ttf"):
            p = Path("/System/Library/Fonts") / n
            if p.exists():
                candidates.append(p)
    else:
        for n in ("NanumGothic.ttf", "NotoSansCJK-Regular.ttc", "malgun.ttf"):
            for base in ("/usr/share/fonts", str(Path.home() / ".fonts")):
                bp = Path(base)
                if not bp.exists():
                    continue
                for p in bp.rglob(n):
                    candidates.append(p)
                    break
    for path in candidates:
        try:
            fm.fontManager.addfont(str(path))
            family = fm.FontProperties(fname=str(path)).get_name()
            matplotlib.rcParams["font.family"] = family
            matplotlib.rcParams["axes.unicode_minus"] = False
            return family, str(path)
        except Exception:
            continue
    matplotlib.rcParams["axes.unicode_minus"] = False
    return None, None


_KO_FONT, _KO_FONT_PATH = _setup_korean_font()


# ─────────────────────────────────────────────────────────────
# Constants
# ─────────────────────────────────────────────────────────────
DIFFICULTY_LEVELS = ["하", "중하", "중", "중상", "상"]
DIFFICULTY_NUM = {"하": 1, "중하": 2, "중": 3, "중상": 4, "상": 5}
VISION_MODEL = "gpt-4o"
TEXT_MODEL = "gpt-4o"
DEFAULT_ACADEMY = "최상위학원"
DEFAULT_PHONE = "0507-1385-4320"

CHART_THEME_OPTIONS = ["editorial", "mono", "vivid"]
CHART_THEME_LABELS = {
    "editorial": "Editorial · 종이 톤",
    "mono":      "Mono · 흑백",
    "vivid":     "Vivid · 활발한 색",
}


# ─────────────────────────────────────────────────────────────
# Schema
# ─────────────────────────────────────────────────────────────
@dataclass
class ExamMeta:
    title: str = ""
    school: str = ""
    grade: str = ""
    subject: str = "영어"
    exam_type: str = "중간고사"
    exam_date: str = ""
    duration_min: int = 50
    total_score: int = 100
    total_questions: int = 0
    notes: str = ""


@dataclass
class Question:
    no: int = 0
    type: str = ""
    difficulty: str = "중"
    score: float = 0.0
    is_subjective: bool = False
    is_killer: bool = False
    scope: str = ""
    memo: str = ""
    passage_excerpt: str = ""
    choices: list[str] = field(default_factory=list)
    # 어려운 문항 자동 크롭에 사용
    page_index: int = 0
    top_ratio: float = 0.0
    bottom_ratio: float = 1.0

    def difficulty_num(self) -> int:
        return DIFFICULTY_NUM.get(self.difficulty, 3)


@dataclass
class KillerDeep:
    """어려운 문항: 시험지 스타일로 재구성한 이미지 + 출제 의도 + 풀이 방법.
    (대비 방법은 보고서 결(結) 섹션에서 통합 제시)
    """
    no: int = 0
    type: str = ""
    score: float = 0.0
    headline: str = ""
    crop_image: bytes = b""             # OCR 텍스트로 재구성한 시험지 스타일 PNG
    page_index: int = 0
    top_ratio: float = 0.0
    bottom_ratio: float = 1.0
    passage_excerpt: str = ""           # 본문 (편집 가능)
    choices: list[str] = field(default_factory=list)  # 선지 (편집 가능)
    intent: str = ""                    # 출제 의도
    solution_method: str = ""           # 풀이 방법


# 블록식 편집기용 — 보고서를 구성하는 단위
@dataclass
class Block:
    id: str = ""
    kind: str = "paragraph"             # heading | paragraph | image | killer | rule
    level: int = 1                      # heading 일 때만 (1~3)
    text: str = ""
    image_bytes: bytes = b""
    caption: str = ""
    killer: dict = field(default_factory=dict)  # KillerDeep 직렬화 (bytes 포함)


def _new_id() -> str:
    return uuid.uuid4().hex[:8]


# ─────────────────────────────────────────────────────────────
# OCR
# ─────────────────────────────────────────────────────────────
OCR_SYSTEM = (
    "당신은 한국 중·고등 영어 시험지를 분석하는 전문가입니다. "
    "이미지에서 시험 메타 정보, 모든 문항의 메타데이터, "
    "그리고 어려운 문항의 깊이 분석에 쓸 지문 핵심 문장과 선지 텍스트를 추출합니다. "
    "또한 각 문항이 시험지 어느 페이지의 어느 세로 영역에 있는지 좌표를 함께 반환합니다. "
    "응답은 반드시 단일 JSON 객체로만 출력하세요."
)

OCR_USER_TEMPLATE = """다음 시험지 이미지(들)을 분석하여 JSON 으로 반환하세요.

[과목 힌트] {subject}
[학교급 힌트] {grade}

[메타정보 추출 — 시험지 표지/머리말/본문에서 가능한 모두 찾아 채우기]
- title: 시험지 상단 제목 그대로 (예: "2025학년도 1학기 1회고사 영어")
- school: 학교명. 헤더/푸터/일련번호 등에서. 예: "이순신고등학교"
- grade: 학년 표시. "1학년", "고2", "중3" 등.
- subject: 과목.
- exam_type: 중간고사/기말고사/모의고사/수행평가/1회고사/2회고사 등.
- exam_date: 시험일자. 표지에 보이면 YYYY-MM-DD 로.
- duration_min, total_score, total_questions: 시험 안내문에서.
- notes: 출제범위 안내문 그대로 옮기기 (예: "출제범위: 1과~2과, 9월 모의고사")

[필수 JSON 스키마]
{{
  "exam_meta": {{
    "title": "...", "school": "...", "grade": "...", "subject": "...",
    "exam_type": "...", "exam_date": "YYYY-MM-DD",
    "duration_min": 50, "total_score": 100, "total_questions": 25,
    "notes": "..."
  }},
  "questions": [
    {{
      "no": 1,
      "type": "어법",
      "difficulty": "중하",
      "score": 4.0,
      "is_subjective": false,
      "scope": "1과",
      "memo": "관계대명사 변형",
      "passage_excerpt": "본문 핵심 문장 1~2개 (정답 단서 부근)",
      "choices": ["① ...", "② ...", "③ ...", "④ ...", "⑤ ..."],
      "page_index": 0,
      "top_ratio": 0.42,
      "bottom_ratio": 0.61
    }}
  ]
}}

[scope (시험범위) — 반드시 짧고 일관된 라벨로]
- 한국 영어 내신은 보통 **교과서 단원** + **모의고사** 출제. 다음 라벨을 우선 사용:
    "1과", "2과", "3과", "4과", "5과", ...
    "모의고사" (어떤 모의 출처든 통일)
    "외부지문" (교과서/모의고사 외)
- 문항 머리/꼬리에 출처 표시(예: "[1과(3)]", "9월 모의 35번", "교과서 Lesson 4")가 있으면 그걸 보고 라벨링.
- 출처 표시가 없으면 본문 내용을 유추해 가까운 단원에 매핑. 모르겠으면 빈 문자열.
- 영문 라벨(Lesson 3) 보다 한국어 라벨(3과) 을 우선.

[좌표 (page_index / top_ratio / bottom_ratio)]
- page_index: 업로드 이미지 인덱스 (0-based).
- top_ratio / bottom_ratio: 그 페이지에서 문항이 차지하는 세로 영역 (0~1).
- 어려운 문항은 **문제 본문 + 정답 근거 지문 핵심부**까지 포함. 약간 여유를 두되 옆 문항이 들어오지 않게.

[난이도 — 정밀 분포 (강조)]
- 상: 시험 전체 1~2문항. 다중 변형·시간 소모·고난도 추론. 매우 엄격하게.
- 중상: 2~4 / 중: 5~8 / 중하: 5~8 / 하: 3~6.
- 모든 문항을 같은 라벨(특히 '중')로 몰지 말 것. 미세 차이를 살려 다섯 단계가 정밀하게 분포하게.

[유형명] 일치/불일치, 빈칸, 어법, 어휘, 순서, 삽입, 요약문, 함의, 대의, 지칭,
영영풀이, 조건영작, 서술형, 단답형 우선 사용.

JSON 외 텍스트는 절대 출력하지 마세요."""


def _img_to_data_url(img_bytes: bytes, mime: str = "image/png") -> str:
    return f"data:{mime};base64,{base64.b64encode(img_bytes).decode('ascii')}"


def _normalize_image(file_bytes: bytes, max_side: int = 2000) -> tuple[bytes, str]:
    try:
        im = Image.open(io.BytesIO(file_bytes)).convert("RGB")
        w, h = im.size
        scale = min(1.0, max_side / max(w, h))
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=88, optimize=True)
        return buf.getvalue(), "image/jpeg"
    except Exception:
        return file_bytes, "image/png"


def _safe_json_loads(text: str) -> dict:
    text = text.strip()
    m = re.search(r"```(?:json)?\s*(.+?)```", text, re.DOTALL)
    if m:
        text = m.group(1).strip()
    return json.loads(text)


def ocr_exam_images(api_key: str, images: list[bytes],
                    subject_hint: str = "영어", grade_hint: str = "") -> tuple[ExamMeta, list[Question]]:
    client = openai.OpenAI(api_key=api_key)
    content: list[dict[str, Any]] = [
        {"type": "text", "text": OCR_USER_TEMPLATE.format(subject=subject_hint, grade=grade_hint)}
    ]
    for raw in images:
        norm, mime = _normalize_image(raw)
        content.append({
            "type": "image_url",
            "image_url": {"url": _img_to_data_url(norm, mime), "detail": "high"},
        })

    resp = client.chat.completions.create(
        model=VISION_MODEL,
        messages=[
            {"role": "system", "content": OCR_SYSTEM},
            {"role": "user", "content": content},
        ],
        temperature=0.1,
        max_tokens=7000,
        response_format={"type": "json_object"},
    )
    data = _safe_json_loads(resp.choices[0].message.content or "{}")

    meta_d = data.get("exam_meta", {}) or {}
    meta = ExamMeta(
        title=str(meta_d.get("title") or "").strip(),
        school=str(meta_d.get("school") or "").strip(),
        grade=str(meta_d.get("grade") or grade_hint).strip(),
        subject=str(meta_d.get("subject") or subject_hint).strip(),
        exam_type=str(meta_d.get("exam_type") or "중간고사").strip(),
        exam_date=str(meta_d.get("exam_date") or "").strip(),
        duration_min=int(meta_d.get("duration_min") or 50),
        total_score=int(meta_d.get("total_score") or 100),
        total_questions=int(meta_d.get("total_questions") or 0),
        notes=str(meta_d.get("notes") or "").strip(),
    )

    qs: list[Question] = []
    for q in data.get("questions", []) or []:
        try:
            diff = q.get("difficulty", "중")
            if diff not in DIFFICULTY_LEVELS:
                diff = "중"
            ch = q.get("choices") or []
            if not isinstance(ch, list):
                ch = []
            top = float(q.get("top_ratio") or 0.0)
            bot = float(q.get("bottom_ratio") or 1.0)
            top = max(0.0, min(1.0, top))
            bot = max(0.0, min(1.0, bot))
            if bot <= top:
                top, bot = 0.0, 1.0
            qs.append(Question(
                no=int(q.get("no", 0) or 0),
                type=str(q.get("type") or "").strip(),
                difficulty=diff,
                score=float(q.get("score") or 0),
                is_subjective=bool(q.get("is_subjective", False)),
                scope=str(q.get("scope") or "").strip(),
                memo=str(q.get("memo") or "").strip(),
                passage_excerpt=str(q.get("passage_excerpt") or "").strip(),
                choices=[str(c).strip() for c in ch if str(c).strip()],
                page_index=int(q.get("page_index") or 0),
                top_ratio=top,
                bottom_ratio=bot,
            ))
        except Exception:
            continue
    qs.sort(key=lambda x: x.no)
    if meta.total_questions == 0:
        meta.total_questions = len(qs)
    return meta, qs


# ─────────────────────────────────────────────────────────────
# Crop helper
# ─────────────────────────────────────────────────────────────
def crop_image_region(image_bytes: bytes, top_ratio: float, bottom_ratio: float,
                       left_ratio: float = 0.0, right_ratio: float = 1.0,
                       pad_ratio: float = 0.01) -> bytes:
    """원본 이미지에서 (top_ratio~bottom_ratio) 세로 영역을 잘라낸 PNG 반환."""
    if not image_bytes:
        return b""
    try:
        im = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    except Exception:
        return b""
    w, h = im.size
    pad = int(h * pad_ratio)
    box = (
        max(0, int(w * left_ratio)),
        max(0, int(h * top_ratio) - pad),
        min(w, int(w * right_ratio)),
        min(h, int(h * bottom_ratio) + pad),
    )
    if box[2] <= box[0] or box[3] <= box[1]:
        return b""
    cropped = im.crop(box)
    out = io.BytesIO()
    cropped.save(out, format="PNG", optimize=True)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────
# 어려운 문항을 시험지 스타일 이미지로 재구성
# ─────────────────────────────────────────────────────────────
def render_question_image(no: int, q_type: str, score: float,
                          passage_excerpt: str, choices: list[str],
                          width: int = 1000) -> bytes:
    """OCR 로 추출한 문항 본문/선지를 깔끔한 시험지 스타일 PNG 로 재구성.
    원본 스캔 크롭이 아니라 타이프셋된 형태."""
    if not (passage_excerpt or "").strip() and not choices:
        return b""

    # 폰트는 _font() 가 PIL 폰트 객체를 반환 (한글 지원)
    F_NUM    = _font(21, bold=True)
    F_META   = _font(13)
    F_BODY   = _font(16)
    F_CHOICE = _font(15)

    PAD = 36
    INNER = width - PAD * 2
    BG = (255, 255, 255)
    INK = (26, 31, 54)
    MUTED = (110, 115, 130)
    BORDER = (212, 207, 192)

    def m_para(text, font, lh=1.7, indent=0) -> int:
        if not text or not text.strip():
            return 0
        return int(font.size * lh) * len(_wrap(text, font, INNER - indent))

    # 캔버스 높이 계산
    H = PAD * 2
    H += int(F_NUM.size * 1.5) + 18              # 헤더
    if (passage_excerpt or "").strip():
        H += m_para(passage_excerpt, F_BODY, indent=12) + 16
    for c in choices:
        H += m_para(c, F_CHOICE, indent=24) + 6

    canvas = Image.new("RGB", (width, max(H, 220)), BG)
    draw = ImageDraw.Draw(canvas)

    # 외곽 테두리 (시험지 느낌)
    draw.rectangle([(1, 1), (width - 2, H - 2)], outline=BORDER, width=1)

    y = PAD
    # 헤더 — 좌: "8. 조건영작" / 우: "5점"
    head_left = f"{no}. {q_type}" if q_type else f"{no}."
    head_right = f"{score:g}점" if score else ""
    draw.text((PAD, y), head_left, font=F_NUM, fill=INK)
    if head_right:
        right_w = F_META.getbbox(head_right)[2]
        draw.text((width - PAD - right_w, y + 6), head_right, font=F_META, fill=MUTED)
    y += int(F_NUM.size * 1.5)
    draw.line([(PAD, y + 4), (width - PAD, y + 4)], fill=BORDER, width=1)
    y += 16

    # 본문
    if (passage_excerpt or "").strip():
        y = _draw_paragraph(draw, passage_excerpt, PAD + 6, y, INNER - 6, F_BODY,
                            color=INK, line_height=1.75)
        y += 14

    # 선지
    for c in choices:
        y = _draw_paragraph(draw, c, PAD + 16, y, INNER - 16, F_CHOICE,
                            color=INK, line_height=1.7)
        y += 4

    final_h = y + PAD
    if final_h < canvas.height:
        canvas = canvas.crop((0, 0, width, final_h))

    out = io.BytesIO()
    canvas.save(out, format="PNG", optimize=True)
    return out.getvalue()


# ─────────────────────────────────────────────────────────────
# 분포 / 라벨
# ─────────────────────────────────────────────────────────────
def auto_killer_flags(meta: ExamMeta, qs: list[Question]) -> list[bool]:
    if not qs:
        return []
    flags = [q.difficulty == "상" for q in qs]
    if not any(flags):
        scores = sorted([q.score for q in qs if q.score > 0])
        threshold = scores[int(len(scores) * 0.75)] if scores else 0
        for i, q in enumerate(qs):
            if q.difficulty == "중상" and q.score >= threshold:
                flags[i] = True
    return flags


def derive_difficulty_label(qs: list[Question]) -> str:
    if not qs:
        return "中"
    nums = [q.difficulty_num() for q in qs]
    avg = sum(nums) / len(nums)
    killer_ratio = sum(1 for q in qs if q.is_killer) / len(qs)
    score = avg + killer_ratio * 1.0
    if score < 1.7:   return "下"
    if score < 2.4:   return "中下"
    if score < 3.1:   return "中"
    if score < 3.8:   return "中上"
    return "上"


def type_distribution(qs: list[Question]) -> list[tuple[str, int, float]]:
    if not qs:
        return []
    s = pd.Series([q.type for q in qs if q.type])
    counts = s.value_counts()
    total = counts.sum()
    return [(k, int(v), v / total * 100) for k, v in counts.items()]


def scope_distribution(qs: list[Question]) -> list[tuple[str, int, float]]:
    s = pd.Series([q.scope for q in qs if q.scope.strip()])
    if s.empty:
        return []
    counts = s.value_counts()
    total = counts.sum()
    return [(k, int(v), v / total * 100) for k, v in counts.items()]


def difficulty_distribution(qs: list[Question]) -> list[tuple[str, int, float]]:
    counts = pd.Series([q.difficulty for q in qs]).value_counts().reindex(DIFFICULTY_LEVELS, fill_value=0)
    total = counts.sum() or 1
    return [(k, int(v), v / total * 100) for k, v in counts.items()]


# ─────────────────────────────────────────────────────────────
# Anti-AI humanize
# ─────────────────────────────────────────────────────────────
AI_TELLS: list[tuple[str, list[str]]] = [
    ("다음과 같이",        ["이렇게", "아래처럼"]),
    ("다음과 같습니다",    ["아래와 같습니다", "정리하면 이렇습니다"]),
    ("종합적으로",        ["전반적으로", "전체를 놓고 보면"]),
    ("효과적으로",        ["제대로", "확실하게"]),
    ("효과적인",          ["제대로 된", "실제로 통하는"]),
    ("매우 ",             ["", "꽤 ", "상당히 "]),
    ("살펴보겠습니다",    ["짚어보면 이렇습니다", "정리해 보면"]),
    ("살펴보면",          ["짚어보면", "들여다보면"]),
    ("이를 통해",         ["이로써", "이렇게 해서"]),
    ("이러한 ",           ["이런 ", ""]),
    ("저희는",            ["", "이번 시험은"]),
    ("여러분",            [""]),
    ("결론적으로",        ["정리하면", "요약하면"]),
    ("나타났습니다",      ["드러났습니다", "확인됐습니다"]),
    ("확인할 수 있습니다", ["확인됩니다", "분명합니다"]),
    ("필요합니다.",        ["필수적입니다.", "요구됩니다."]),
    ("중요합니다.",        ["관건입니다.", "핵심입니다."]),
    ("출제되었습니다",    ["출제됐습니다", "출제된 셈입니다"]),
]
OPENING_VARIANTS = [
    "이번 시험은",
    "전반적으로 보면 이번 시험은",
    "결과만 놓고 보면",
    "체감 난이도부터 짚어보면",
    "큰 틀에서 이번 시험은",
]


def humanize_text(text: str, seed: int | None = None) -> str:
    if not text:
        return ""
    rng = random.Random(seed)
    out = text
    for ai_phrase, alts in AI_TELLS:
        if ai_phrase in out:
            out = out.replace(ai_phrase, rng.choice(alts), 1)
    sentences = re.split(r"(?<=[.!?])\s+", out)
    if len(sentences) >= 3:
        ips = [i for i, s in enumerate(sentences) if s.endswith("입니다.")]
        if ips and rng.random() < 0.25:
            i = rng.choice(ips)
            sentences[i] = sentences[i][:-4] + "라 할 수 있습니다."
        out = " ".join(sentences)
    out = re.sub(r"분석한 결과,?\s*", "", out)
    out = re.sub(r"^\s*결과적으로,?\s*", "", out, flags=re.MULTILINE)
    out = re.sub(r"[ \t]+", " ", out)
    out = re.sub(r"\n{3,}", "\n\n", out)
    return out.strip()


# ─────────────────────────────────────────────────────────────
# 어려운 문항 깊이 분석 — 함정 / 풀이 / 대비 (페러프레이징은 이미지가 대신함)
# ─────────────────────────────────────────────────────────────
KILLER_DEEP_SYSTEM = """당신은 입시 영어 분석 전문가입니다.
어려운 문항 각각에 대해, 이미지로 보여줄 문제는 그대로 두고
다음 2가지만 학부모와 학생이 이해할 수 있는 한국어로 작성합니다:

1. intent (출제 의도)
   왜 이 문항을 이렇게 냈는지. 어떤 능력을 변별하려는 의도인지.
   왜 학생들이 어려워하는지(어디서 막히고, 어떤 함정이 박혀 있는지)까지 한 단락에 자연스럽게 담아주세요.
   2~3 문장.

2. solution_method (풀이 방법)
   접근 순서. 어떤 단서를 먼저 잡고 어떤 순으로 좁혀가야 하는지 구체적인 행동 지침.
   학생이 시험장에서 쓸 수 있는 실용 절차로 써주세요. 2~3 문장.

추가:
- headline: 한 줄로 이 문항의 본질 요약 (예: "본문 변형 + 어법 조건 결합형")

[금지 표현]
"다음과 같이", "다양한", "효과적으로", "이를 통해", "여러분", "체계적으로",
"~을 통해서", "결론적으로 말씀드리면", "활용하여"

문체: 학원 원장이 학부모에게 차분히 설명하는 단정한 한국어. "~합니다" 종결.
※ 대비 방법은 별도 섹션에서 따로 다루므로 여기서는 다루지 마세요.
JSON 으로만 응답하세요."""


def gen_killer_deep(api_key: str, meta: ExamMeta, qs: list[Question],
                    original_images: list[bytes]) -> list[KillerDeep]:
    killers = [q for q in qs if q.is_killer]
    if not killers:
        return []
    client = openai.OpenAI(api_key=api_key)

    items_input = [{
        "no": q.no, "type": q.type, "difficulty": q.difficulty, "score": q.score,
        "is_subjective": q.is_subjective, "scope": q.scope, "memo": q.memo,
        "passage_excerpt": q.passage_excerpt, "choices": q.choices,
    } for q in killers]

    user = f"""[시험 정보]
{meta.school} {meta.grade} {meta.subject} {meta.exam_type}
총 {meta.total_questions}문항 · {meta.total_score}점 · {meta.duration_min}분

[어려운 문항 데이터]
{json.dumps(items_input, ensure_ascii=False, indent=2)}

요청: 각 문항에 대해 출제 의도 + 풀이 방법 + 한 줄 헤드라인을 작성하세요.

응답 JSON:
{{
  "items": [
    {{
      "no": 4,
      "type": "조건영작",
      "headline": "본문 변형 + 어법 조건 결합형",
      "intent": "...",
      "solution_method": "..."
    }}
  ]
}}

JSON 외 텍스트 없이 반환하세요."""

    resp = client.chat.completions.create(
        model=TEXT_MODEL,
        messages=[
            {"role": "system", "content": KILLER_DEEP_SYSTEM},
            {"role": "user", "content": user},
        ],
        temperature=0.55,
        max_tokens=2000,
        response_format={"type": "json_object"},
    )
    data = _safe_json_loads(resp.choices[0].message.content or "{}")

    by_no = {q.no: q for q in killers}
    out: list[KillerDeep] = []
    for it in data.get("items", []) or []:
        no = int(it.get("no", 0) or 0)
        q = by_no.get(no)
        if q is None:
            continue
        # OCR 텍스트를 시험지 스타일로 재구성 (스캔 크롭 대신)
        img_b = render_question_image(
            no=q.no, q_type=q.type, score=q.score,
            passage_excerpt=q.passage_excerpt, choices=q.choices,
        )
        # 재구성이 비어있으면(passage/choices 둘 다 없음) 원본 크롭으로 폴백
        if not img_b and 0 <= q.page_index < len(original_images):
            img_b = crop_image_region(
                original_images[q.page_index],
                q.top_ratio, q.bottom_ratio,
            )
        out.append(KillerDeep(
            no=no,
            type=str(it.get("type", q.type)),
            score=q.score,
            headline=humanize_text(it.get("headline", "")),
            crop_image=img_b,
            page_index=q.page_index,
            top_ratio=q.top_ratio,
            bottom_ratio=q.bottom_ratio,
            passage_excerpt=q.passage_excerpt,
            choices=list(q.choices),
            intent=humanize_text(it.get("intent", "")),
            solution_method=humanize_text(it.get("solution_method", "")),
        ))
    return out


# ─────────────────────────────────────────────────────────────
# 본문 — 기 / 승 / 결 + 차트 캡션 (단일 호출, JSON)
# ─────────────────────────────────────────────────────────────
LEEPIN_SAMPLE = """이번 시험은 객관적 난이도 '중하' 수준으로, 기본기가 충실한 학생이라면 무난히 풀 수 있는 구성이었습니다. 다만 일부 문항의 선택지가 교묘하게 짜여 체감 난이도는 '중'까지 올라갔을 것으로 보입니다. 서술형 난도가 하락한 점은 주목할 변화이며, 객관식에서의 한 번 실수가 등급을 가르는 구조였습니다.

균형 잡힌 범위 분포가 이번 시험의 큰 그림입니다. 5과·6과가 36%씩 동등하게 출제되었고 모의고사가 28%를 차지해, 단원 편식을 허용하지 않는 구성이었습니다. 변별의 분기점은 '일치/불일치'와 '어휘'에 있었습니다. 두 유형이 전체의 44%를 차지하며 꼼꼼한 지문 분석 능력을 직접 평가했습니다. 서답형은 평이해 점수를 얻는 구간이 아닌, 실점하지 않고 지켜내야 하는 구간으로 작용했습니다.

다음 시험을 위한 학습 전략은 분명합니다. 첫째, 전 범위 균형 학습이 필수적입니다. 한 단원에만 시간을 쏟는 방식으로는 상위 등급이 어렵습니다. 둘째, 지문 완전 정복에 집중해야 합니다. 단순 해석을 넘어 세부 정보까지 근거를 잡아 푸는 훈련이 요구됩니다. 셋째, 서답형 실수 방지 루틴이 필요합니다. 핵심 구문과 어휘는 의미뿐 아니라 정확한 철자까지 완벽하게 쓰는 연습을 매일 반복하셔야 합니다. 평이한 기조에 안주하기보다, 기본기를 단단히 다지면서 고난도 변화에도 흔들리지 않는 실력을 만드는 것이 장기적 전략입니다."""


BLOG_BODY_SYSTEM = """당신은 학원 원장 LEEPIN 입니다. 영어 입시 전문, 중·고등 내신을 직접 분석해 블로그에 글을 쓰는 사람입니다.

[글쓰기 스타일]
- 학부모와 학생이 함께 읽는 블로그 톤. 단단한 문어체이지만 "갈렸습니다", "치명적인", "변별의 분기점", "기회이자 위기" 같은 입시 지도자 표현 자연스럽게 사용
- 객관 데이터와 전문가 판단을 한 단락에 섞기
- 짧은 문장과 긴 문장을 섞어 리듬 만들기
- 이모지 금지

[금지 표현]
"다음과 같이", "분석한 결과", "살펴보겠습니다", "여러분", "효과적으로",
"종합적으로", "다양한", "이를 통해", "이러한 측면에서", "결론적으로 말씀드리면",
"~을 통해서", "체계적으로", "활용하여",
"~인 점은 인상적입니다", "~을 알 수 있습니다", "~로 사료됩니다"

[보고서 흐름은 기승전결 — 당신은 그 중 세 부분(gi/seung/gyeol)을 씁니다]
  · 기 = 총평 (gi)
  · 승 = 시험 기조 (seung) + 차트 4종과 캡션
  · 전 = 어려운 문항 풀이 (코드가 별도 카드로 삽입, 당신이 쓰지 않음)
  · 결 = 대비 방법 (gyeol)

[작성 가이드 — 틀에 갇히지 말고 시험 성격에 맞게 유연하게]
- 기 (gi): 시험 한눈 진단. 객관/체감 난이도, 출제 기조 한 마디, 다음 시험 시사.
  분량은 시험 성격에 따라 3~6문장 사이에서 자연스럽게.
- 승 (seung): 시험 기조 — 유형·범위·난이도 분포가 어떤 학습 신호를 주는지.
  데이터는 본문에 녹여 쓰기. 표/불릿 만들지 말 것. 핵심 신호 2~4개 정도.
- 결 (gyeol): 다음 시험을 위한 대비 방법.
  실용적 행동 지침을 적으세요. 분량/빈도/대상이 손에 잡혀야 합니다.
  형식은 자유롭게 — 첫째·둘째·셋째 형태도 좋고, 자연스러운 단락 전개도 좋습니다.
  시험이 단순했다면 짧고 단단하게, 복잡했다면 단계적으로 나눠 써도 됩니다.
  마지막은 장기 학습 태도로 단단하게 닫습니다.

[차트 캡션 4종]
각 차트 아래 1~3 문장. 데이터의 의미를 짧게 짚기.

전부 한 번에 JSON 으로 응답."""


def gen_blog_body(api_key: str, meta: ExamMeta, qs: list[Question],
                  killer_deeps: list[KillerDeep], academy: str, phone: str,
                  seed: int | None = None) -> dict:
    if not qs:
        return {"gi": "", "seung": "", "gyeol": "", "captions": {}}
    rng = random.Random(seed)
    client = openai.OpenAI(api_key=api_key)

    diff_label = derive_difficulty_label(qs)
    type_dist = type_distribution(qs)
    scope_dist = scope_distribution(qs)
    diff_dist = difficulty_distribution(qs)

    title = (meta.title or " ".join(filter(None, [
        meta.school, meta.grade, meta.exam_date[:7] if meta.exam_date else "",
        meta.exam_type, "영어 시험 분석",
    ])))

    type_lines = "\n".join(f"- {t}: {n}문항 ({p:.1f}%)" for t, n, p in type_dist)
    scope_lines = "\n".join(f"- {s}: {n}문항 ({p:.1f}%)" for s, n, p in scope_dist) or "(범위 정보 미제공)"
    diff_lines = "\n".join(f"- {d}: {n}문항 ({p:.1f}%)" for d, n, p in diff_dist if n > 0)
    killer_brief = "\n".join(f"- {kd.no}번 ({kd.type}): {kd.headline}" for kd in killer_deeps) or "- 명시적 어려운 문항 없음"

    user = f"""[참고: 당신이 과거에 쓴 글의 톤]
{LEEPIN_SAMPLE}

────────────────
[이번 시험]
제목: {title}
난도: {diff_label} · 총 {meta.total_questions}문항 · {meta.total_score}점 · {meta.duration_min}분

[유형 분포]
{type_lines}

[범위 분포]
{scope_lines}

[난이도 분포]
{diff_lines}

[어려운 문항 헤드라인 (전(轉) 영역에 별도 카드로 들어감)]
{killer_brief}

────────────────
지시:
- 시작 어구 예: "{rng.choice(OPENING_VARIANTS)}" (그대로 쓰지 말고 자연스럽게 변형)
- 표·불릿 나열은 만들지 마세요. 데이터는 본문에 녹여 쓰기.
- 결(gyeol) 마지막 단락은 학습 태도/대비 방법으로 강하게 마무리.
- 차트 캡션 4종은 각 1~3 문장. {('범위 데이터가 없으므로 scope 캡션은 빈 문자열' if not scope_dist else '')}

응답:
{{
  "gi": "...",
  "seung": "...",
  "gyeol": "...",
  "captions": {{"type":"...","scope":"...","difficulty":"...","location":"..."}}
}}
"""

    draft = client.chat.completions.create(
        model=TEXT_MODEL,
        messages=[
            {"role": "system", "content": BLOG_BODY_SYSTEM},
            {"role": "user", "content": user},
        ],
        temperature=0.85,
        max_tokens=2400,
        response_format={"type": "json_object"},
    ).choices[0].message.content or "{}"

    polish_sys = (
        "당신은 한국 입시 학원 원장의 글을 자연스럽게 다듬는 편집자입니다. "
        "AI 가 쓴 듯한 매끄러움을 줄이고 사람이 손으로 쓴 호흡을 살리세요. "
        "금지 표현 '다음과 같이', '분석한 결과', '살펴보겠습니다', '효과적으로', '종합적으로', "
        "'다양한', '이를 통해', '체계적으로', '활용하여' 를 자연스러운 표현으로 모두 대체하세요. "
        "원문 JSON 키 구조와 의미는 보존하고, 같은 JSON 으로 반환하세요."
    )
    polished = client.chat.completions.create(
        model=TEXT_MODEL,
        messages=[
            {"role": "system", "content": polish_sys},
            {"role": "user", "content": draft},
        ],
        temperature=0.6,
        max_tokens=2400,
        response_format={"type": "json_object"},
    ).choices[0].message.content or draft

    data = _safe_json_loads(polished)
    captions = data.get("captions", {}) or {}
    return {
        "gi": humanize_text(data.get("gi", ""), seed=seed),
        "seung": humanize_text(data.get("seung", ""), seed=(seed or 0) + 1),
        "gyeol": humanize_text(data.get("gyeol", ""), seed=(seed or 0) + 2),
        "captions": {
            "type": humanize_text(captions.get("type", "")),
            "scope": humanize_text(captions.get("scope", "")),
            "difficulty": humanize_text(captions.get("difficulty", "")),
            "location": humanize_text(captions.get("location", "")),
        },
    }


# ─────────────────────────────────────────────────────────────
# 차트 — 테마별 팔레트
# ─────────────────────────────────────────────────────────────
def _editorial_style(ax, palette):
    ax.set_facecolor(palette["paper"])
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    for s in ("left", "bottom"):
        ax.spines[s].set_color(palette["rule"])
        ax.spines[s].set_linewidth(0.8)
    ax.tick_params(colors=palette["ink"], labelsize=10)
    ax.grid(axis="y", color=palette["rule"], linestyle="-", linewidth=0.5, alpha=0.6)
    ax.set_axisbelow(True)


def chart_type_distribution(qs: list[Question], theme: ThemeName = "editorial") -> bytes:
    pal = chart_palette(theme)
    dist = type_distribution(qs)
    if not dist:
        return b""
    labels, counts, pcts = zip(*dist)
    fig, ax = plt.subplots(figsize=(9, max(4.5, 0.55 * len(labels))), facecolor=pal["paper"])
    bars = ax.barh(labels[::-1], counts[::-1],
                   color=[pal["accents"][i % len(pal["accents"])] for i in range(len(labels))][::-1],
                   edgecolor=pal["paper"], height=0.7)
    for bar, n, p in zip(bars, counts[::-1], pcts[::-1]):
        ax.text(n + max(counts) * 0.02, bar.get_y() + bar.get_height() / 2,
                f"{n}문항 ({p:.1f}%)", va="center", fontsize=10.5, color=pal["ink"])
    ax.set_xlim(0, max(counts) * 1.32)
    ax.set_title("유형별 출제 비중", loc="left", fontsize=14, color=pal["ink"], fontweight="bold", pad=14)
    _editorial_style(ax, pal)
    ax.tick_params(axis="x", labelsize=0)
    ax.spines["bottom"].set_visible(False)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, facecolor=pal["paper"], bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def chart_scope_distribution(qs: list[Question], theme: ThemeName = "editorial") -> bytes:
    pal = chart_palette(theme)
    dist = scope_distribution(qs)
    if not dist:
        return b""
    labels, counts, pcts = zip(*dist)
    fig, ax = plt.subplots(figsize=(8, max(3.5, 0.55 * len(labels))), facecolor=pal["paper"])
    bars = ax.barh(labels[::-1], counts[::-1],
                   color=[pal["accents"][(i + 2) % len(pal["accents"])] for i in range(len(labels))][::-1],
                   edgecolor=pal["paper"], height=0.7)
    for bar, n, p in zip(bars, counts[::-1], pcts[::-1]):
        ax.text(n + max(counts) * 0.02, bar.get_y() + bar.get_height() / 2,
                f"{n}문항 ({p:.1f}%)", va="center", fontsize=10.5, color=pal["ink"])
    ax.set_xlim(0, max(counts) * 1.32)
    ax.set_title("범위별 문항 분포", loc="left", fontsize=14, color=pal["ink"], fontweight="bold", pad=14)
    _editorial_style(ax, pal)
    ax.tick_params(axis="x", labelsize=0)
    ax.spines["bottom"].set_visible(False)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, facecolor=pal["paper"], bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def chart_difficulty_distribution(qs: list[Question], theme: ThemeName = "editorial") -> bytes:
    pal = chart_palette(theme)
    dist = difficulty_distribution(qs)
    fig, ax = plt.subplots(figsize=(8, 3.6), facecolor=pal["paper"])
    labels, counts, pcts = zip(*dist)
    bars = ax.bar(labels, counts,
                  color=[pal["accents"][i % len(pal["accents"])] for i in range(len(labels))],
                  edgecolor=pal["paper"], linewidth=2, width=0.55)
    for bar, n, p in zip(bars, counts, pcts):
        if n > 0:
            ax.text(bar.get_x() + bar.get_width() / 2, n + max(counts) * 0.02,
                    f"{n}\n({p:.1f}%)", ha="center", fontsize=9.5,
                    color=pal["ink"], fontweight="bold")
    ax.set_title("난이도 분포", loc="left", fontsize=14, color=pal["ink"], fontweight="bold", pad=14)
    _editorial_style(ax, pal)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, facecolor=pal["paper"], bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


def chart_killer_map(qs: list[Question], theme: ThemeName = "editorial") -> bytes:
    pal = chart_palette(theme)
    if not qs:
        return b""
    fig, ax = plt.subplots(figsize=(9, 3.0), facecolor=pal["paper"])
    nos = [q.no for q in qs]
    ds = [q.difficulty_num() for q in qs]
    killer_color = pal["accents"][3] if len(pal["accents"]) > 3 else pal["accents"][-1]
    colors = [killer_color if q.is_killer else pal["ink"] for q in qs]
    sizes = [120 if q.is_killer else 36 for q in qs]
    ax.scatter(nos, ds, c=colors, s=sizes, edgecolor=pal["paper"], linewidth=1.2, zorder=3)
    ax.plot(nos, ds, color=pal["ink"], alpha=0.16, linewidth=1, zorder=1)
    for q in qs:
        if q.is_killer:
            ax.annotate(f"#{q.no}", (q.no, q.difficulty_num()),
                        xytext=(0, 12), textcoords="offset points",
                        ha="center", fontsize=9.5, color=killer_color, fontweight="bold")
    ax.set_yticks(list(DIFFICULTY_NUM.values()))
    ax.set_yticklabels(list(DIFFICULTY_NUM.keys()))
    ax.set_xlabel("문항 번호", fontsize=10, color=pal["ink"])
    ax.set_title("문항 위치별 난이도 & 어려운 문항 분포", loc="left",
                 fontsize=14, color=pal["ink"], fontweight="bold", pad=14)
    _editorial_style(ax, pal)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=200, facecolor=pal["paper"], bbox_inches="tight")
    plt.close(fig)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────
# Block 빌드 / 직렬화 헬퍼
# ─────────────────────────────────────────────────────────────
def _killer_to_dict(kd: KillerDeep) -> dict:
    return {
        "no": kd.no, "type": kd.type, "score": kd.score, "headline": kd.headline,
        "crop_image": kd.crop_image,
        "page_index": kd.page_index,
        "top_ratio": kd.top_ratio, "bottom_ratio": kd.bottom_ratio,
        "passage_excerpt": kd.passage_excerpt,
        "choices": list(kd.choices),
        "intent": kd.intent,
        "solution_method": kd.solution_method,
    }


def body_to_blocks(meta: ExamMeta, qs: list[Question],
                   gi: str, seung: str, gyeol: str, captions: dict,
                   killer_deeps: list[KillerDeep], charts: dict,
                   academy: str, phone: str) -> list[Block]:
    diff_label = derive_difficulty_label(qs)
    blocks: list[Block] = []

    # 기 — 총평
    blocks.append(Block(id=_new_id(), kind="heading", level=1, text=f"총평 (난도: {diff_label})"))
    blocks.append(Block(id=_new_id(), kind="paragraph", text=gi))

    # 승 — 시험 기조 + 차트들 (이미지 → 캡션)
    blocks.append(Block(id=_new_id(), kind="heading", level=1, text="시험 기조"))
    if seung.strip():
        blocks.append(Block(id=_new_id(), kind="paragraph", text=seung))
    chart_specs = [
        ("type", "유형별 출제 비중", "type"),
        ("scope", "범위별 문항 분포", "scope"),
        ("difficulty", "난이도 분포", "difficulty"),
        ("killer_map", "문항 위치별 난이도 & 어려운 문항", "location"),
    ]
    for key, _, cap_key in chart_specs:
        if charts.get(key):
            blocks.append(Block(
                id=_new_id(), kind="image",
                image_bytes=charts[key],
                caption=(captions or {}).get(cap_key, ""),
            ))

    # 전 — 어려운 문항 풀이
    if killer_deeps:
        blocks.append(Block(id=_new_id(), kind="heading", level=1,
                            text="어려운 문항 풀이"))
        for kd in killer_deeps:
            blocks.append(Block(id=_new_id(), kind="killer", killer=_killer_to_dict(kd)))

    # 결 — 대비 방법
    blocks.append(Block(id=_new_id(), kind="heading", level=1, text="대비 방법"))
    if gyeol.strip():
        blocks.append(Block(id=_new_id(), kind="paragraph", text=gyeol))

    # 푸터
    blocks.append(Block(id=_new_id(), kind="paragraph",
                        text=f"수강 문의: {academy} ☎️ {phone}"))
    return blocks


# ─────────────────────────────────────────────────────────────
# PIL utilities
# ─────────────────────────────────────────────────────────────
def _font(size: int, bold: bool = False) -> ImageFont.ImageFont:
    path = _KO_FONT_PATH
    bold_path = None
    if path:
        p = Path(path)
        for cand in (p.parent / "malgunbd.ttf", p.parent / "NanumGothicBold.ttf"):
            if cand.exists():
                bold_path = str(cand)
                break
    use = (bold_path if bold and bold_path else path)
    if use:
        try:
            return ImageFont.truetype(use, size)
        except Exception:
            pass
    return ImageFont.load_default()


def _wrap(text: str, font: ImageFont.ImageFont, max_w: int) -> list[str]:
    out = []
    for paragraph in text.split("\n"):
        if not paragraph.strip():
            out.append("")
            continue
        line = ""
        for ch in paragraph:
            test = line + ch
            w = font.getbbox(test)[2]
            if w > max_w and line:
                out.append(line)
                line = ch
            else:
                line = test
        if line:
            out.append(line)
    return out


def _draw_paragraph(draw: ImageDraw.ImageDraw, text: str, x: int, y: int, max_w: int,
                    font: ImageFont.ImageFont, color, line_height: float = 1.65) -> int:
    if not text or not text.strip():
        return y
    lines = _wrap(text, font, max_w)
    line_h = int(font.size * line_height)
    for ln in lines:
        draw.text((x, y), ln, font=font, fill=color)
        y += line_h
    return y


# ─────────────────────────────────────────────────────────────
# 블록 → PNG / Word
# ─────────────────────────────────────────────────────────────
INK = (26, 31, 54)
MUTED = (110, 115, 130)
ACCENT = (45, 58, 92)
RULE = (212, 207, 192)


def render_blog_image_from_blocks(meta: ExamMeta, qs: list[Question],
                                   blocks: list[Block]) -> bytes:
    W = 1080
    PAD = 60
    INNER = W - PAD * 2
    BG = (255, 255, 255)

    F_TITLE = _font(38, bold=True)
    F_H1    = _font(26, bold=True)
    F_H2    = _font(20, bold=True)
    F_H3    = _font(17, bold=True)
    F_BODY  = _font(17)
    F_SMALL = _font(15)
    F_META  = _font(14)

    def m_para(text, font, lh=1.65) -> int:
        if not text or not text.strip():
            return 0
        return int(font.size * lh) * len(_wrap(text, font, INNER))

    # 측정 — 캔버스 높이 계산
    H = PAD * 2
    H += int(F_TITLE.size * 1.4) + 30
    H += int(F_META.size * 1.6) + 28 + 1 + 24
    for blk in blocks:
        if blk.kind == "heading":
            font = {1: F_H1, 2: F_H2, 3: F_H3}.get(blk.level, F_H1)
            H += int(font.size * 1.4) * len(_wrap(blk.text, font, INNER)) + 16
        elif blk.kind == "paragraph":
            H += m_para(blk.text, F_BODY) + 18
        elif blk.kind == "image":
            try:
                im = Image.open(io.BytesIO(blk.image_bytes))
                ratio = INNER / im.width
                H += int(im.height * ratio) + 12
            except Exception:
                pass
            if blk.caption.strip():
                H += m_para(blk.caption, F_SMALL) + 12
        elif blk.kind == "killer":
            kd = blk.killer
            H += int(F_H2.size * 1.4) + 8
            if kd.get("crop_image"):
                try:
                    im = Image.open(io.BytesIO(kd["crop_image"]))
                    ratio = INNER / im.width
                    H += int(im.height * ratio) + 14
                except Exception:
                    pass
            for k in ("trap_analysis", "solution_method", "prep_method"):
                if kd.get(k, "").strip():
                    H += int(F_H3.size * 1.4) + m_para(kd[k], F_BODY) + 14
            H += 12

    canvas = Image.new("RGB", (W, max(H, 1200)), BG)
    draw = ImageDraw.Draw(canvas)
    y = PAD

    # 타이틀
    title = (meta.title or " ".join(filter(None, [
        meta.school, meta.grade, meta.exam_date[:7] if meta.exam_date else "",
        meta.exam_type, "영어 시험 분석",
    ])))
    for ln in _wrap(title, F_TITLE, INNER):
        draw.text((PAD, y), ln, font=F_TITLE, fill=INK)
        y += int(F_TITLE.size * 1.2)
    y += 10

    diff_label = derive_difficulty_label(qs)
    meta_str = " · ".join(filter(None, [
        meta.school or "", meta.grade or "", f"총 {meta.total_questions}문항",
        f"{meta.total_score}점", f"{meta.duration_min}분", f"난도 {diff_label}",
    ]))
    draw.text((PAD, y), meta_str, font=F_META, fill=MUTED)
    y += int(F_META.size * 1.4) + 18
    draw.line([(PAD, y), (W - PAD, y)], fill=ACCENT, width=2)
    y += 26

    KILLER_LABELS = [
        ("intent",          "출제 의도"),
        ("solution_method", "풀이 방법"),
    ]

    for blk in blocks:
        if blk.kind == "heading":
            y += 10
            font = {1: F_H1, 2: F_H2, 3: F_H3}.get(blk.level, F_H1)
            for ln in _wrap(blk.text, font, INNER):
                draw.text((PAD, y), ln, font=font, fill=INK)
                y += int(font.size * 1.3)
            y += 6
        elif blk.kind == "paragraph":
            y = _draw_paragraph(draw, blk.text, PAD, y, INNER, F_BODY, color=INK)
            y += 14
        elif blk.kind == "image":
            try:
                im = Image.open(io.BytesIO(blk.image_bytes)).convert("RGB")
                ratio = INNER / im.width
                new_h = int(im.height * ratio)
                im = im.resize((INNER, new_h), Image.LANCZOS)
                canvas.paste(im, (PAD, y))
                y += new_h + 8
            except Exception:
                pass
            if blk.caption.strip():
                y = _draw_paragraph(draw, blk.caption, PAD, y, INNER, F_SMALL, color=MUTED, line_height=1.6)
                y += 10
        elif blk.kind == "killer":
            kd = blk.killer
            head = f"{kd.get('no','')}번 · {kd.get('type','')} — {kd.get('headline','')}"
            for ln in _wrap(head, F_H2, INNER):
                draw.text((PAD, y), ln, font=F_H2, fill=ACCENT)
                y += int(F_H2.size * 1.25)
            y += 6
            if kd.get("crop_image"):
                try:
                    im = Image.open(io.BytesIO(kd["crop_image"])).convert("RGB")
                    ratio = INNER / im.width
                    new_h = int(im.height * ratio)
                    im = im.resize((INNER, new_h), Image.LANCZOS)
                    canvas.paste(im, (PAD, y))
                    y += new_h + 12
                except Exception:
                    pass
            for k, lbl in KILLER_LABELS:
                v = kd.get(k, "")
                if not v.strip():
                    continue
                draw.text((PAD, y), lbl, font=F_H3, fill=ACCENT)
                y += int(F_H3.size * 1.4)
                y = _draw_paragraph(draw, v, PAD + 18, y, INNER - 18, F_BODY, color=INK)
                y += 8
            y += 14

    # 자르기
    final_h = y + PAD
    if final_h < canvas.height:
        canvas = canvas.crop((0, 0, W, final_h))
    out = io.BytesIO()
    canvas.save(out, format="PNG", optimize=True)
    return out.getvalue()


# ── Word ──
def _set_east_asia(run, name="맑은 고딕"):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def _add_run(p, text, *, size=10.5, bold=False, color="1A1F36", mono=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    if mono:
        r.font.name = "IBM Plex Mono"
    else:
        r.font.name = "맑은 고딕"
        _set_east_asia(r)
    return r


def build_word_report_from_blocks(meta: ExamMeta, qs: list[Question],
                                   blocks: list[Block]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = Cm(1.6)
    section.left_margin = section.right_margin = Cm(1.8)

    diff_label = derive_difficulty_label(qs)

    title = (meta.title or " ".join(filter(None, [
        meta.school, meta.grade, meta.exam_date[:7] if meta.exam_date else "",
        meta.exam_type, "영어 시험 분석",
    ])))
    p = doc.add_paragraph()
    _add_run(p, title, size=22, bold=True, color="1A1F36")
    p = doc.add_paragraph()
    meta_str = " · ".join(filter(None, [
        meta.school or "", meta.grade or "",
        f"총 {meta.total_questions}문항", f"{meta.total_score}점", f"{meta.duration_min}분",
        f"난도 {diff_label}",
    ]))
    _add_run(p, meta_str, size=10, color="6E7382")
    doc.add_paragraph()

    KILLER_LABELS = [
        ("intent",          "출제 의도"),
        ("solution_method", "풀이 방법"),
    ]

    for blk in blocks:
        if blk.kind == "heading":
            doc.add_paragraph()
            p = doc.add_paragraph()
            size = {1: 14, 2: 12.5, 3: 11.5}.get(blk.level, 14)
            _add_run(p, blk.text, size=size, bold=True, color="1A1F36")
        elif blk.kind == "paragraph":
            for para in blk.text.split("\n\n"):
                para = para.strip()
                if not para:
                    continue
                p = doc.add_paragraph()
                _add_run(p, para, size=10.5, color="1A1F36")
                p.paragraph_format.line_spacing = 1.65
                p.paragraph_format.space_after = Pt(8)
        elif blk.kind == "image":
            if blk.image_bytes:
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p_img.add_run().add_picture(io.BytesIO(blk.image_bytes), width=Cm(16))
                except Exception:
                    pass
            if blk.caption.strip():
                p = doc.add_paragraph()
                _add_run(p, blk.caption, size=10, color="6E7382")
                p.paragraph_format.line_spacing = 1.6
        elif blk.kind == "killer":
            kd = blk.killer
            p = doc.add_paragraph()
            _add_run(p, f"{kd.get('no','')}번 · {kd.get('type','')} ", size=12, bold=True, color="2D3A5C", mono=True)
            _add_run(p, f"— {kd.get('headline','')}", size=12, bold=True, color="2D3A5C")
            if kd.get("crop_image"):
                p_img = doc.add_paragraph()
                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p_img.add_run().add_picture(io.BytesIO(kd["crop_image"]), width=Cm(15))
                except Exception:
                    pass
            for k, lbl in KILLER_LABELS:
                v = kd.get(k, "")
                if not v.strip():
                    continue
                p = doc.add_paragraph()
                _add_run(p, lbl, size=10.5, bold=True, color="2D3A5C")
                p = doc.add_paragraph()
                _add_run(p, v, size=10.5, color="1A1F36")
                p.paragraph_format.line_spacing = 1.6
                p.paragraph_format.left_indent = Cm(0.5)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_blog_text_from_blocks(meta: ExamMeta, qs: list[Question],
                                 blocks: list[Block]) -> str:
    diff_label = derive_difficulty_label(qs)
    title = (meta.title or " ".join(filter(None, [
        meta.school, meta.grade, meta.exam_date[:7] if meta.exam_date else "",
        meta.exam_type, "영어 시험 분석",
    ])))
    out = [title, ""]
    out.append(f"({meta.school or ''} {meta.grade or ''} · 총 {meta.total_questions}문항 · {meta.total_score}점 · 난도 {diff_label})")
    out.append("")
    for blk in blocks:
        if blk.kind == "heading":
            out.append(blk.text)
            out.append("")
        elif blk.kind == "paragraph":
            out.append(blk.text)
            out.append("")
        elif blk.kind == "image":
            out.append("[이미지]")
            if blk.caption.strip():
                out.append(blk.caption)
            out.append("")
        elif blk.kind == "killer":
            kd = blk.killer
            out.append(f"{kd.get('no','')}번 · {kd.get('type','')} — {kd.get('headline','')}")
            out.append("")
            if kd.get("crop_image"):
                out.append("[문제 + 정답 근거 이미지]")
                out.append("")
            for k, lbl in [("intent", "출제 의도"),
                           ("solution_method", "풀이 방법")]:
                v = kd.get(k, "")
                if v.strip():
                    out.append(lbl)
                    out.append(v)
                    out.append("")
    return "\n".join(out)


# ─────────────────────────────────────────────────────────────
# UI
# ─────────────────────────────────────────────────────────────
SS_PREFIX = "exam_"


def _ss(key: str, default=None):
    full = SS_PREFIX + key
    if full not in st.session_state:
        st.session_state[full] = default
    return st.session_state[full]


def _set_ss(key: str, value):
    st.session_state[SS_PREFIX + key] = value


def _init_state():
    _ss("meta", None)
    _ss("questions", [])
    _ss("killer_deeps", [])
    _ss("blocks", [])
    _ss("blog_text", "")
    _ss("blog_image", b"")
    _ss("blog_word", b"")
    _ss("uploaded_keys", [])
    _ss("original_images", [])
    _ss("academy", DEFAULT_ACADEMY)
    _ss("phone", DEFAULT_PHONE)
    _ss("chart_theme", "editorial")
    _ss("editing_block_id", None)


def render_sidebar():
    _init_state()

    st.markdown("### 분석 설정")
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    st.markdown(
        "<p class='section-label'>OCR 설정</p>"
        "<div style='font-size:13px;color:var(--text-body);line-height:1.7'>"
        "과목 <span style='color:var(--text-muted)'>·</span> "
        "<b>영어</b> <span class='tag tag-neutral' style='margin-left:4px'>고정</span><br/>"
        "학교/학년 <span style='color:var(--text-muted)'>·</span> "
        "<span style='color:var(--text-muted)'>OCR 자동 추출</span>"
        "</div>",
        unsafe_allow_html=True,
    )

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<p class="section-label">차트 테마</p>', unsafe_allow_html=True)
    st.caption("보고서 차트 4종의 색감을 정합니다.")
    cur = _ss("chart_theme", "editorial")
    new_ct = st.radio(
        "chart_theme_radio", CHART_THEME_OPTIONS,
        index=CHART_THEME_OPTIONS.index(cur) if cur in CHART_THEME_OPTIONS else 0,
        format_func=lambda k: CHART_THEME_LABELS[k],
        label_visibility="collapsed", key="chart_theme_radio",
    )
    if new_ct != cur:
        _set_ss("chart_theme", new_ct)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<p class="section-label">학원 정보 (보고서 푸터)</p>', unsafe_allow_html=True)
    st.text_input("학원명", key=SS_PREFIX + "academy", value=_ss("academy", DEFAULT_ACADEMY))
    st.text_input("연락처", key=SS_PREFIX + "phone", value=_ss("phone", DEFAULT_PHONE))

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<p class="section-label">초기화</p>', unsafe_allow_html=True)
    if st.button("분석 결과 초기화", use_container_width=True):
        for k in ("meta", "questions", "killer_deeps", "blocks",
                  "blog_text", "blog_image", "blog_word",
                  "uploaded_keys", "original_images"):
            if k == "meta":
                st.session_state[SS_PREFIX + k] = None
            elif k in ("questions", "killer_deeps", "blocks", "uploaded_keys", "original_images"):
                st.session_state[SS_PREFIX + k] = []
            elif k in ("blog_image", "blog_word"):
                st.session_state[SS_PREFIX + k] = b""
            else:
                st.session_state[SS_PREFIX + k] = ""
        _set_ss("editing_block_id", None)
        st.rerun()

    meta: ExamMeta | None = _ss("meta")
    if meta is not None:
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown('<p class="section-label">현재 시험</p>', unsafe_allow_html=True)
        st.markdown(
            f"<div style='font-size:13px;color:var(--text-body);line-height:1.6'>"
            f"<b>{meta.title or '제목 없음'}</b><br/>"
            f"{meta.school or ''} {meta.grade or ''} {meta.subject or ''}<br/>"
            f"<span style='color:var(--text-muted)'>"
            f"{meta.exam_type} · {meta.total_questions}문항 · {meta.total_score}점"
            f"</span></div>",
            unsafe_allow_html=True,
        )


def _meta_card(meta: ExamMeta, qs: list[Question]) -> str:
    killer_count = sum(1 for q in qs if q.is_killer)
    return f"""
<div class='kpi-strip'>
  <div class='card'>
    <div class='card-eyebrow'>총 문항</div>
    <div class='card-value'>{meta.total_questions}</div>
    <div class='card-meta'>{meta.subject} · {meta.exam_type}</div>
  </div>
  <div class='card'>
    <div class='card-eyebrow'>총 배점</div>
    <div class='card-value'>{meta.total_score}<span style='font-size:14px;color:var(--text-muted)'> 점</span></div>
    <div class='card-meta'>{meta.duration_min}분</div>
  </div>
  <div class='card'>
    <div class='card-eyebrow'>학교 · 학년</div>
    <div class='card-value' style='font-size:18px'>{meta.grade or '—'}</div>
    <div class='card-meta'>{meta.school or '—'}</div>
  </div>
  <div class='card'>
    <div class='card-eyebrow'>어려운 문항</div>
    <div class='card-value' style='color:var(--killer-fg)'>{killer_count}</div>
    <div class='card-meta'>난도 {derive_difficulty_label(qs)}</div>
  </div>
</div>
"""


def _questions_to_df(qs: list[Question]) -> pd.DataFrame:
    if not qs:
        return pd.DataFrame(columns=["no", "type", "difficulty", "score",
                                     "is_subjective", "is_killer", "scope", "memo"])
    rows = [{
        "no": q.no, "type": q.type, "difficulty": q.difficulty, "score": q.score,
        "is_subjective": q.is_subjective, "is_killer": q.is_killer,
        "scope": q.scope, "memo": q.memo,
    } for q in qs]
    return pd.DataFrame(rows)


def _df_to_questions(df: pd.DataFrame, prev_qs: list[Question]) -> list[Question]:
    by_no = {q.no: q for q in prev_qs}
    qs = []
    for _, row in df.iterrows():
        try:
            no = int(row.get("no") or 0)
            base = by_no.get(no, Question(no=no))
            base.no = no
            base.type = str(row.get("type") or "").strip()
            base.difficulty = str(row.get("difficulty") or "중")
            base.score = float(row.get("score") or 0)
            base.is_subjective = bool(row.get("is_subjective") or False)
            base.is_killer = bool(row.get("is_killer") or False)
            base.scope = str(row.get("scope") or "").strip()
            base.memo = str(row.get("memo") or "").strip()
            qs.append(base)
        except Exception:
            continue
    qs.sort(key=lambda x: x.no)
    return qs


# ─────────────────────────────────────────────────────────────
# Block 편집 헬퍼
# ─────────────────────────────────────────────────────────────
def _blocks() -> list[Block]:
    return _ss("blocks", [])


def _move_block(idx: int, delta: int):
    bs = _blocks()
    new_idx = idx + delta
    if 0 <= new_idx < len(bs):
        bs[idx], bs[new_idx] = bs[new_idx], bs[idx]
        _set_ss("blocks", bs)


def _delete_block(idx: int):
    bs = _blocks()
    if 0 <= idx < len(bs):
        del bs[idx]
        _set_ss("blocks", bs)


def _insert_block(idx: int, kind: str):
    bs = _blocks()
    new_blk = Block(id=_new_id(), kind=kind)
    if kind == "heading":
        new_blk.level = 2
        new_blk.text = "새 제목"
    elif kind == "paragraph":
        new_blk.text = "새 본문 단락"
    bs.insert(idx, new_blk)
    _set_ss("blocks", bs)


def _block_summary(blk: Block) -> str:
    if blk.kind == "heading":
        return f"H{blk.level} · {blk.text[:60]}"
    if blk.kind == "paragraph":
        t = blk.text.replace("\n", " ")
        return f"본문 · {t[:80]}{'…' if len(t) > 80 else ''}"
    if blk.kind == "image":
        cap = f" — {blk.caption[:40]}" if blk.caption else ""
        return f"이미지{cap}"
    if blk.kind == "killer":
        kd = blk.killer
        return f"어려운 문항 · #{kd.get('no','')} {kd.get('type','')}"
    return blk.kind


def _render_block_editor():
    """블록 리스트를 보여주고 각 블록의 수정/이동/삭제 UI 제공."""
    bs = _blocks()
    if not bs:
        st.markdown("<div class='empty'>아직 보고서가 없습니다. 위에서 [블로그용 분석 보고서 생성] 을 눌러주세요.</div>",
                    unsafe_allow_html=True)
        return

    editing_id = _ss("editing_block_id")

    for idx, blk in enumerate(bs):
        is_editing = (editing_id == blk.id)
        with st.container(border=True):
            c1, c2, c3, c4, c5 = st.columns([7, 1, 1, 1, 1])
            with c1:
                st.markdown(
                    f"<div style='font-size:13px;color:var(--text-faint);font-family:var(--font-mono);text-transform:uppercase;letter-spacing:.06em'>{blk.kind}</div>"
                    f"<div style='font-size:14px;color:var(--text-body);margin-top:2px'>{_block_summary(blk)}</div>",
                    unsafe_allow_html=True,
                )
            with c2:
                if st.button("✎", key=f"edit_{blk.id}", help="수정", use_container_width=True):
                    _set_ss("editing_block_id", None if is_editing else blk.id)
                    st.rerun()
            with c3:
                if st.button("↑", key=f"up_{blk.id}", help="위로", use_container_width=True,
                             disabled=(idx == 0)):
                    _move_block(idx, -1)
                    st.rerun()
            with c4:
                if st.button("↓", key=f"down_{blk.id}", help="아래로", use_container_width=True,
                             disabled=(idx == len(bs) - 1)):
                    _move_block(idx, +1)
                    st.rerun()
            with c5:
                if st.button("✕", key=f"del_{blk.id}", help="삭제", use_container_width=True):
                    _delete_block(idx)
                    if editing_id == blk.id:
                        _set_ss("editing_block_id", None)
                    st.rerun()

            # 인라인 편집 폼
            if is_editing:
                _render_inline_editor(blk, idx)

        # 블록 사이에 추가 버튼
        ic1, ic2, _ = st.columns([1, 1, 6])
        with ic1:
            if st.button("＋ 텍스트", key=f"ins_t_{blk.id}", help="이 아래에 텍스트 블록 추가"):
                _insert_block(idx + 1, "paragraph")
                st.rerun()
        with ic2:
            if st.button("＋ 이미지", key=f"ins_i_{blk.id}", help="이 아래에 이미지 블록 추가"):
                _insert_block(idx + 1, "image")
                _set_ss("editing_block_id", _blocks()[idx + 1].id)
                st.rerun()


def _render_inline_editor(blk: Block, idx: int):
    bs = _blocks()
    st.markdown('<div style="margin-top:8px"></div>', unsafe_allow_html=True)
    if blk.kind == "heading":
        new_text = st.text_input("제목 텍스트", blk.text, key=f"h_text_{blk.id}")
        new_level = st.selectbox("크기", [1, 2, 3], index=[1, 2, 3].index(blk.level) if blk.level in (1, 2, 3) else 0,
                                 format_func=lambda x: f"H{x}", key=f"h_lvl_{blk.id}")
        if st.button("저장", key=f"save_{blk.id}", type="primary"):
            blk.text = new_text
            blk.level = new_level
            _set_ss("blocks", bs)
            _set_ss("editing_block_id", None)
            st.rerun()
    elif blk.kind == "paragraph":
        new_text = st.text_area("본문 텍스트", blk.text, height=180, key=f"p_text_{blk.id}")
        if st.button("저장", key=f"save_{blk.id}", type="primary"):
            blk.text = new_text
            _set_ss("blocks", bs)
            _set_ss("editing_block_id", None)
            st.rerun()
    elif blk.kind == "image":
        if blk.image_bytes:
            st.image(blk.image_bytes, caption=blk.caption or None, use_container_width=True)
        new_caption = st.text_input("캡션 (선택)", blk.caption, key=f"i_cap_{blk.id}")
        new_file = st.file_uploader("이미지 파일 (PNG/JPG)", type=["png", "jpg", "jpeg", "webp"],
                                     key=f"i_file_{blk.id}")
        c1, c2 = st.columns(2)
        with c1:
            if st.button("저장", key=f"save_{blk.id}", type="primary", use_container_width=True):
                if new_file is not None:
                    new_file.seek(0)
                    blk.image_bytes = new_file.read()
                blk.caption = new_caption
                _set_ss("blocks", bs)
                _set_ss("editing_block_id", None)
                st.rerun()
        with c2:
            if st.button("취소", key=f"cancel_{blk.id}", use_container_width=True):
                _set_ss("editing_block_id", None)
                st.rerun()
    elif blk.kind == "killer":
        kd = blk.killer
        st.markdown(f"**#{kd.get('no')} · {kd.get('type','')}**")
        new_head = st.text_input("헤드라인", kd.get("headline", ""), key=f"k_head_{blk.id}")

        # 현재 이미지 미리보기
        if kd.get("crop_image"):
            st.image(kd["crop_image"], caption="시험지 스타일로 재구성한 문항 이미지", use_container_width=True)

        # 본문 / 선지 편집 + 이미지 재생성
        with st.expander("문항 본문/선지 편집", expanded=False):
            new_passage = st.text_area("본문 (passage_excerpt)", kd.get("passage_excerpt", ""),
                                       height=160, key=f"k_passage_{blk.id}",
                                       help="시험지 스타일 이미지로 재구성될 본문 텍스트입니다.")
            choices_text = "\n".join(kd.get("choices", []))
            new_choices_text = st.text_area("선지 (한 줄에 하나씩, 빈 줄 무시)",
                                            choices_text, height=140, key=f"k_choices_{blk.id}",
                                            help="예시:\n① first\n② second\n③ third")
            if st.button("이미지 재생성", key=f"k_regen_{blk.id}"):
                cs = [c.strip() for c in new_choices_text.split("\n") if c.strip()]
                kd["passage_excerpt"] = new_passage
                kd["choices"] = cs
                kd["crop_image"] = render_question_image(
                    no=int(kd.get("no", 0)), q_type=kd.get("type", ""),
                    score=float(kd.get("score", 0)),
                    passage_excerpt=new_passage, choices=cs,
                )
                _set_ss("blocks", bs)
                st.rerun()

        # 원본 스캔 크롭으로 대체하고 싶을 때
        originals: list[bytes] = _ss("original_images", [])
        page_idx = int(kd.get("page_index", 0))
        if 0 <= page_idx < len(originals):
            with st.expander("원본 스캔으로 대체 (필요할 때만)", expanded=False):
                top = st.slider("상단 위치", 0.0, 1.0, float(kd.get("top_ratio", 0.0)),
                                step=0.005, key=f"k_top_{blk.id}", format="%.3f")
                bot = st.slider("하단 위치", 0.0, 1.0, float(kd.get("bottom_ratio", 1.0)),
                                step=0.005, key=f"k_bot_{blk.id}", format="%.3f")
                if st.button("원본 스캔 크롭으로 교체", key=f"k_recrop_{blk.id}"):
                    if bot <= top:
                        st.warning("하단이 상단보다 작거나 같습니다.")
                    else:
                        kd["top_ratio"] = top
                        kd["bottom_ratio"] = bot
                        kd["crop_image"] = crop_image_region(originals[page_idx], top, bot)
                        _set_ss("blocks", bs)
                        st.rerun()

        new_intent = st.text_area("출제 의도", kd.get("intent", ""), height=120, key=f"k_intent_{blk.id}")
        new_sol = st.text_area("풀이 방법", kd.get("solution_method", ""), height=120, key=f"k_sol_{blk.id}")
        if st.button("저장", key=f"save_{blk.id}", type="primary"):
            kd["headline"] = new_head
            kd["intent"] = new_intent
            kd["solution_method"] = new_sol
            _set_ss("blocks", bs)
            _set_ss("editing_block_id", None)
            st.rerun()


# ─────────────────────────────────────────────────────────────
# 메인
# ─────────────────────────────────────────────────────────────
def render_main(api_key: str):
    if not api_key:
        st.error("OpenAI API Key 가 필요합니다.")
        return
    _init_state()
    chart_theme: ThemeName = _ss("chart_theme", "editorial")

    # ── §1 ──
    st.markdown('<div class="section-mark">§ 1. 업로드 & OCR</div>', unsafe_allow_html=True)

    files = st.file_uploader(
        "시험지 이미지 업로드 (여러 페이지 동시 가능)",
        type=["png", "jpg", "jpeg", "webp"],
        accept_multiple_files=True, key="exam_uploader",
    )

    if files:
        keys = [f"{f.name}_{f.size}" for f in files]
        if keys != _ss("uploaded_keys"):
            _set_ss("uploaded_keys", keys)
        thumb_cols = st.columns(min(len(files), 4))
        for i, f in enumerate(files):
            with thumb_cols[i % 4]:
                st.image(f, use_container_width=True, caption=f.name)

        if st.button("OCR 실행 — 메타정보 & 문항 추출", type="primary", use_container_width=True):
            with st.status("이미지 분석 중...", expanded=True) as status:
                try:
                    img_bytes_list = []
                    for f in files:
                        f.seek(0)
                        img_bytes_list.append(f.read())
                    _set_ss("original_images", img_bytes_list)
                    status.update(label=f"GPT-4o Vision 으로 {len(files)}장 분석 중...")
                    meta, qs = ocr_exam_images(api_key, img_bytes_list, "영어", "")
                    flags = auto_killer_flags(meta, qs)
                    for q, flag in zip(qs, flags):
                        q.is_killer = flag
                    _set_ss("meta", meta)
                    _set_ss("questions", qs)
                    _set_ss("killer_deeps", [])
                    _set_ss("blocks", [])
                    _set_ss("blog_text", "")
                    _set_ss("blog_image", b"")
                    _set_ss("blog_word", b"")
                    status.update(label=f"완료 — 문항 {len(qs)}개 추출", state="complete")
                except openai.AuthenticationError:
                    status.update(label="API Key 오류", state="error")
                    st.error("OpenAI API Key 가 유효하지 않습니다.")
                except Exception as e:
                    status.update(label="OCR 실패", state="error")
                    st.error(f"분석 중 오류: {e}")

    meta: ExamMeta | None = _ss("meta")
    qs: list[Question] = _ss("questions")
    if meta is None:
        st.markdown('<div class="empty">이미지를 업로드하고 OCR 을 실행하면 여기에 분석 결과가 표시됩니다.</div>',
                    unsafe_allow_html=True)
        return

    # ── §2 ──
    st.markdown('<div class="section-mark" style="margin-top:32px">§ 2. 메타정보 확인 & 수정</div>',
                unsafe_allow_html=True)
    st.markdown(_meta_card(meta, qs), unsafe_allow_html=True)

    with st.expander("시험 메타정보 수정", expanded=False):
        c1, c2, c3 = st.columns(3)
        with c1:
            meta.title = st.text_input("시험 제목", meta.title, key="m_title")
            meta.school = st.text_input("학교명", meta.school, key="m_school")
            meta.grade = st.text_input("학년", meta.grade, key="m_grade")
        with c2:
            meta.subject = st.text_input("과목", meta.subject, key="m_subject")
            options = ["중간고사", "기말고사", "모의고사", "수행평가", "기타"]
            meta.exam_type = st.selectbox("시험 종류", options,
                                          index=options.index(meta.exam_type) if meta.exam_type in options else 0,
                                          key="m_type")
            meta.exam_date = st.text_input("시험일자", meta.exam_date, key="m_date")
        with c3:
            meta.duration_min = st.number_input("시험 시간 (분)", 10, 200, meta.duration_min, key="m_dur")
            meta.total_score = st.number_input("총 배점", 10, 200, meta.total_score, key="m_tot")
            meta.notes = st.text_input("출제 범위/메모", meta.notes, key="m_notes")
        _set_ss("meta", meta)

    st.markdown("##### 문항 정보 — 표에서 직접 수정")
    df = _questions_to_df(qs)
    edited = st.data_editor(
        df,
        column_config={
            "no": st.column_config.NumberColumn("번호", width=60, format="%d"),
            "type": st.column_config.TextColumn("유형", width=100),
            "difficulty": st.column_config.SelectboxColumn("난이도", options=DIFFICULTY_LEVELS, width=80),
            "score": st.column_config.NumberColumn("배점", min_value=0, max_value=20, step=0.5, width=70),
            "is_subjective": st.column_config.CheckboxColumn("서답형", width=70),
            "is_killer": st.column_config.CheckboxColumn("어려운 문항", width=90),
            "scope": st.column_config.TextColumn("범위", width=110),
            "memo": st.column_config.TextColumn("메모", width=200),
        },
        use_container_width=True, hide_index=True, num_rows="dynamic",
        key="exam_q_editor",
    )
    if not edited.equals(df):
        _set_ss("questions", _df_to_questions(edited, qs))
        qs = _ss("questions")

    qa, qb, qc = st.columns(3)
    with qa:
        if st.button("어려운 문항 자동 표시", use_container_width=True):
            flags = auto_killer_flags(meta, qs)
            for q, f in zip(qs, flags):
                q.is_killer = f
            _set_ss("questions", qs)
            st.rerun()
    with qb:
        if st.button("어려운 문항 모두 해제", use_container_width=True):
            for q in qs:
                q.is_killer = False
            _set_ss("questions", qs)
            st.rerun()
    with qc:
        st.download_button(
            "메타정보 JSON 저장",
            data=json.dumps({"meta": asdict(meta),
                             "questions": [asdict(q) for q in qs]},
                            ensure_ascii=False, indent=2).encode("utf-8"),
            file_name=f"exam_meta_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            mime="application/json",
            use_container_width=True,
        )

    # ── §3 — 보고서 생성 ──
    st.markdown('<div class="section-mark" style="margin-top:32px">§ 3. 블로그용 보고서 생성</div>',
                unsafe_allow_html=True)
    st.caption("이 버튼 한 번이면 — 어려운 문항 자동 크롭 + 함정/풀이/대비 + 기·승·결 + 차트까지 모두 만들어 §4에서 블록식으로 편집 가능합니다.")

    academy = _ss("academy", DEFAULT_ACADEMY)
    phone = _ss("phone", DEFAULT_PHONE)

    if st.button("블로그용 분석 보고서 생성", type="primary", use_container_width=True):
        if not qs:
            st.warning("문항 정보가 비어 있습니다.")
        else:
            try:
                with st.status("어려운 문항 깊이 분석 + 자동 크롭 중...", expanded=True) as status:
                    originals: list[bytes] = _ss("original_images", [])
                    kds = gen_killer_deep(api_key, meta, qs, originals)
                    _set_ss("killer_deeps", kds)

                    status.update(label="기·승·결 본문 작성 중...")
                    body = gen_blog_body(api_key, meta, qs, kds, academy, phone,
                                         seed=hash((meta.title, meta.exam_date)) & 0xFFFFFFFF)

                    status.update(label="차트 생성 중...")
                    charts = {
                        "type":       chart_type_distribution(qs, chart_theme),
                        "scope":      chart_scope_distribution(qs, chart_theme),
                        "difficulty": chart_difficulty_distribution(qs, chart_theme),
                        "killer_map": chart_killer_map(qs, chart_theme),
                    }

                    status.update(label="블록 구성 중...")
                    blocks = body_to_blocks(meta, qs, body["gi"], body["seung"], body["gyeol"],
                                            body["captions"], kds, charts, academy, phone)
                    _set_ss("blocks", blocks)

                    status.update(label="초기 산출물 생성 중...")
                    word_b = build_word_report_from_blocks(meta, qs, blocks)
                    img_b = render_blog_image_from_blocks(meta, qs, blocks)
                    text_b = build_blog_text_from_blocks(meta, qs, blocks)
                    _set_ss("blog_word", word_b)
                    _set_ss("blog_image", img_b)
                    _set_ss("blog_text", text_b)
                    status.update(label="완료 — §4 에서 편집할 수 있습니다", state="complete")
            except openai.AuthenticationError:
                st.error("OpenAI API Key 가 유효하지 않습니다.")
            except Exception as e:
                st.error(f"분석 실패: {e}")

    # ── §4 — 블록 편집기 ──
    if _blocks():
        st.markdown('<div class="section-mark" style="margin-top:32px">§ 4. 블록식 편집</div>',
                    unsafe_allow_html=True)
        st.caption("각 블록을 ✎(수정) / ↑↓(이동) / ✕(삭제) 하거나 사이에 ＋ 로 새 블록을 끼워 넣을 수 있습니다. 어려운 문항 블록은 크롭 영역도 미세조정 가능합니다.")

        ac1, ac2 = st.columns([2, 1])
        with ac1:
            if st.button("변경사항 적용 → PNG · Word 재생성", type="primary", use_container_width=True):
                blocks = _blocks()
                _set_ss("blog_word", build_word_report_from_blocks(meta, qs, blocks))
                _set_ss("blog_image", render_blog_image_from_blocks(meta, qs, blocks))
                _set_ss("blog_text", build_blog_text_from_blocks(meta, qs, blocks))
                st.success("산출물이 갱신되었습니다.")
        with ac2:
            if st.button("맨 위에 + 텍스트", use_container_width=True):
                _insert_block(0, "paragraph")
                st.rerun()

        _render_block_editor()

        # 맨 아래 추가
        st.markdown("##### 블록 추가")
        bc1, bc2 = st.columns(2)
        with bc1:
            if st.button("＋ 텍스트 블록을 맨 아래에", use_container_width=True):
                _insert_block(len(_blocks()), "paragraph")
                st.rerun()
        with bc2:
            if st.button("＋ 이미지 블록을 맨 아래에", use_container_width=True):
                _insert_block(len(_blocks()), "image")
                _set_ss("editing_block_id", _blocks()[-1].id)
                st.rerun()

    # ── §5 — 결과 미리보기 + 다운로드 ──
    blog_image = _ss("blog_image", b"")
    blog_word = _ss("blog_word", b"")
    blog_text = _ss("blog_text", "")
    if blog_image or blog_text:
        st.markdown('<div class="section-mark" style="margin-top:32px">§ 5. 결과 미리보기 & 다운로드</div>',
                    unsafe_allow_html=True)
        prev_col, dl_col = st.columns([3, 2])
        with prev_col:
            if blog_image:
                st.markdown("##### 블로그 이미지 미리보기")
                st.image(blog_image, use_container_width=True)
        with dl_col:
            st.markdown("##### 다운로드")
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            base = re.sub(r'[\\/*?:"<>|]', "_",
                          meta.title or f"{meta.school}_{meta.grade}_{meta.subject}_{meta.exam_type}")
            if blog_word:
                st.download_button(
                    "Word 보고서 (.docx)", data=blog_word,
                    file_name=f"{base}_분석_{ts}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary", use_container_width=True,
                )
            if blog_image:
                st.download_button(
                    "블로그 이미지 (.png)", data=blog_image,
                    file_name=f"{base}_분석_{ts}.png", mime="image/png",
                    type="primary", use_container_width=True,
                )
            if blog_text:
                st.download_button(
                    "블로그 텍스트 (.txt)", data=blog_text.encode("utf-8"),
                    file_name=f"{base}_분석_{ts}.txt", mime="text/plain",
                    use_container_width=True,
                )
