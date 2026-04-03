"""
최상위학원 영어 정기시험 보고서 자동화 시스템
Streamlit Web App
"""

import streamlit as st
import pandas as pd
import os
import re
import io
import zipfile
import tempfile
import platform
import subprocess
import random
from datetime import datetime
from statistics import mean
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
import openai


# ════════════════════════════════════════════════════════════
# 0. FONT INSTALL (Linux 환경에서 맑은 고딕 설치)
# ════════════════════════════════════════════════════════════
if platform.system() != "Windows":
    _font_dir = Path(__file__).parent / ".fonts"
    _sys_font_dir = Path.home() / ".fonts"
    if _font_dir.exists() and not (_sys_font_dir / "malgun.ttf").exists():
        _sys_font_dir.mkdir(exist_ok=True)
        import shutil
        for f in _font_dir.glob("*.ttf"):
            shutil.copy2(f, _sys_font_dir / f.name)
        subprocess.run(["fc-cache", "-fv"], capture_output=True)

# ════════════════════════════════════════════════════════════
# 1. PAGE CONFIG
# ════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="최상위학원 보고서 자동화",
    page_icon="",
    layout="wide",
    initial_sidebar_state="expanded",
)


# ════════════════════════════════════════════════════════════
# 2. API KEY (자동 로드 — 선생님은 몰라도 됨)
# ════════════════════════════════════════════════════════════
@st.cache_data
def load_api_key():
    # 1) Streamlit secrets (Cloud 배포용)
    try:
        key = st.secrets.get("OPENAI_API_KEY", "")
        if key:
            return key
    except Exception:
        pass

    # 2) 환경 변수
    key = os.environ.get("OPENAI_API_KEY", "")
    if key:
        return key

    # 3) .env 파일 탐색
    base = Path(__file__).parent
    candidates = [
        base / ".env", base / ".env.txt", base / "env.txt",
        base.parent / ".env", base.parent / ".env.txt",
    ]
    for p in candidates:
        if p.exists():
            try:
                for line in p.read_text(encoding="utf-8").splitlines():
                    line = line.strip()
                    if line and not line.startswith("#") and "=" in line:
                        k, v = line.split("=", 1)
                        k = k.strip().lstrip("\ufeff")
                        v = v.strip().strip('"').strip("'")
                        if k == "OPENAI_API_KEY" and v:
                            return v
            except Exception:
                pass
    return ""

API_KEY = load_api_key()

@st.cache_data
def load_cloudconvert_key():
    try:
        key = st.secrets.get("CLOUDCONVERT_API_KEY", "")
        if key:
            return key
    except Exception:
        pass
    return os.environ.get("CLOUDCONVERT_API_KEY", "")

CC_API_KEY = load_cloudconvert_key()

# 고정 AI 설정
AI_MODEL = "gpt-5.1"
AI_TEMPERATURE = 0.7
AI_MAX_TOKENS = 500


# ════════════════════════════════════════════════════════════
# 3. DESIGN SYSTEM — CSS Custom Properties
# ════════════════════════════════════════════════════════════
st.markdown("""
<style>
@import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard@v1.3.9/dist/web/variable/pretendardvariable-dynamic-subset.min.css');

:root {
    --c-primary:     #4A5A8C;
    --c-primary-50:  #F0F2F8;
    --c-primary-100: #DDE1EF;
    --c-primary-600: #3D4D7A;
    --c-gray-50:  #F9FAFB;
    --c-gray-100: #F3F4F6;
    --c-gray-200: #E5E7EB;
    --c-gray-300: #D1D5DB;
    --c-gray-400: #9CA3AF;
    --c-gray-500: #6B7280;
    --c-gray-700: #374151;
    --c-gray-900: #111827;
    --c-success:    #059669;
    --c-success-bg: #ECFDF5;
    --c-warn:       #D97706;
    --c-warn-bg:    #FFFBEB;
    --c-error:      #DC2626;
    --c-error-bg:   #FEF2F2;
    --radius-sm: 6px;
    --radius-md: 10px;
    --radius-lg: 14px;
    --shadow-sm: 0 1px 3px rgba(0,0,0,0.06), 0 1px 2px rgba(0,0,0,0.04);
    --shadow-md: 0 4px 6px -1px rgba(0,0,0,0.06), 0 2px 4px -1px rgba(0,0,0,0.04);
}

/* ── Global ── */
html, body, .stApp,
.stApp [class*="css"] {
    font-family: 'Pretendard Variable', 'Pretendard', -apple-system,
                 BlinkMacSystemFont, 'Malgun Gothic', sans-serif !important;
}
.stApp { background: var(--c-gray-50); }

/* ── Sidebar ── */
section[data-testid="stSidebar"] {
    background: #FFFFFF;
    border-right: 1px solid var(--c-gray-200);
}
section[data-testid="stSidebar"] .stMarkdown p {
    font-size: 13px;
    color: var(--c-gray-500);
}

/* ── Typography ── */
h1 {
    color: var(--c-primary) !important;
    font-weight: 700 !important;
    font-size: 1.75rem !important;
    letter-spacing: -0.5px;
}
h2, h3 {
    color: var(--c-gray-900) !important;
    font-weight: 600 !important;
}
h3 { font-size: 1.05rem !important; }

/* ── Cards ── */
.card {
    background: #fff;
    border: 1px solid var(--c-gray-200);
    border-radius: var(--radius-lg);
    padding: 20px 22px;
    margin-bottom: 10px;
    transition: box-shadow .15s ease;
}
.card:hover { box-shadow: var(--shadow-md); }
.card-label {
    font-size: 11px;
    font-weight: 600;
    color: var(--c-gray-400);
    text-transform: uppercase;
    letter-spacing: .6px;
    margin-bottom: 6px;
}
.card-value {
    font-size: 26px;
    font-weight: 700;
    color: var(--c-primary);
    line-height: 1.2;
}
.card-meta {
    font-size: 12.5px;
    color: var(--c-gray-500);
    margin-top: 6px;
}

/* ── Tag / Badge ── */
.tag {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 9999px;
    font-size: 11px;
    font-weight: 600;
    letter-spacing: .2px;
}
.tag-done   { background: var(--c-success-bg); color: var(--c-success); }
.tag-partial { background: var(--c-warn-bg);    color: var(--c-warn); }
.tag-empty  { background: var(--c-primary-50); color: var(--c-primary); }

/* ── Buttons ── */
.stButton > button {
    border-radius: var(--radius-md) !important;
    font-weight: 500 !important;
    padding: .5rem 1.25rem !important;
    transition: all .12s ease !important;
    font-size: 14px !important;
}

/* ── Progress ── */
.stProgress > div > div > div {
    background: linear-gradient(90deg, var(--c-primary), var(--c-primary-600)) !important;
    border-radius: 9999px !important;
}

/* ── Tabs ── */
.stTabs [data-baseweb="tab-list"] {
    gap: 0;
    border-bottom: 2px solid var(--c-gray-200);
}
.stTabs [data-baseweb="tab"] {
    border-radius: 0 !important;
    padding: 10px 28px !important;
    font-weight: 500 !important;
    font-size: 14px !important;
    color: var(--c-gray-500) !important;
    border-bottom: 2px solid transparent;
    margin-bottom: -2px;
}
.stTabs [aria-selected="true"] {
    color: var(--c-primary) !important;
    border-bottom-color: var(--c-primary) !important;
    background: transparent !important;
}

/* ── Divider ── */
.divider {
    height: 1px;
    background: var(--c-gray-200);
    margin: 20px 0;
}
.divider-accent {
    height: 2px;
    background: linear-gradient(90deg, var(--c-primary), transparent);
    margin: 12px 0 24px 0;
    border-radius: 1px;
}

/* ── Empty state ── */
.empty {
    text-align: center;
    padding: 56px 24px;
    color: var(--c-gray-400);
    font-size: 14px;
}

/* ── Section label (sidebar) ── */
.section-label {
    font-size: 10px;
    font-weight: 700;
    color: var(--c-gray-400);
    text-transform: uppercase;
    letter-spacing: 1.2px;
    margin: 24px 0 8px 0;
}

/* ── Hide defaults ── */
#MainMenu, footer { visibility: hidden; }
.stDeployButton { display: none; }
</style>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════
# 4. DEFAULT PROMPTS
# ════════════════════════════════════════════════════════════
DEFAULT_SYSTEM_PROMPT = (
    "당신은 영어 입시 전문학원의 담당 선생님으로, 학부모에게 보내는 시험 보고서 코멘트를 작성합니다. "
    "매번 다른 문장 구조와 표현을 사용하여, 학생마다 고유한 느낌의 코멘트를 만들어야 합니다. "
    "정형화된 틀이나 상투적 표현을 피하고, 실제 선생님이 직접 쓴 것처럼 자연스러운 문체를 사용하세요."
)

DEFAULT_USER_PROMPT = """아래 학생의 시험 결과를 보고, 학부모님께 전달할 코멘트를 작성해 주세요.

[학생 정보]
- 이름: {student_name}
- Reading: {reading_score}점 (반 평균 {reading_avg}점)
- Grammar: {grammar_score}점 (반 평균 {grammar_avg}점)
- 수업 태도: {attitude}
- 과제 성실성: {sincerity}
- 특이사항: {notes}
{keyword_section}
[작성 규칙]
1. 인사말 없이 바로 시작. 5~8문장, 한 단락.
2. 매번 다른 방식으로 시작하세요. "~학생은"으로 시작하지 마세요. 수업 중 에피소드, 점수 언급, 태도 묘사 등 다양하게.
3. 점수를 반 평균과 비교하되, "반 평균과 동일한 수준을 보여 주었는데"같은 틀에 박힌 표현 대신 자연스럽게 녹이세요.
4. 잘하는 부분은 구체적으로, 부족한 부분은 "~하면 좋겠습니다" 식으로 부드럽게.
5. 특이사항에 적힌 내용을 반드시 반영하되, 그대로 옮기지 말고 자연스럽게 풀어쓰세요.
{keyword_guideline}
6. "인상적입니다", "돋보입니다", "기대됩니다" 같은 표현은 코멘트당 최대 1회만. 같은 표현이 여러 학생에게 반복되지 않도록 매번 다른 방식으로 표현하세요.
7. 문장 길이를 짧은 것과 긴 것을 섞어 리듬감 있게.
8. **반드시 한국어로 작성.**"""

ALL_CLASSES = ["월수_1부", "월수_2부", "월수_3부", "화목_1부", "화목_2부", "화목_3부"]


FORTUNES = [
    "학부모가 감동하여 간식을 보내올 기운이 느껴집니다.",
    "오늘 퇴근은 빠를 예정입니다. 아마도.",
    "학생들이 숙제를 다 해올 확률: 3.7%",
    "커피 한 잔의 여유가 필요한 시점입니다.",
    "오늘의 행운의 단어: comprehensive",
    "학생이 문법 만점을 받으면 로또를 사세요.",
    "오늘따라 학생들이 조용할 운세입니다. 믿거나 말거나.",
    "반 평균이 5점 오르면 치킨 각입니다.",
    "Reading 지문이 술술 읽히는 하루가 될 것입니다. 학생 말고 선생님이요.",
    "오늘 수업 중 졸는 학생 수: 평소보다 1명 적음.",
    "시험지 출력하다 잉크 안 떨어지면 좋은 하루입니다.",
    "오늘의 행운의 문법: 가정법 과거완료",
    "어디선가 학생이 선생님 흉내를 내고 있습니다. 좋은 뜻으로요.",
    "오늘은 빨간 펜보다 파란 펜이 행운을 부릅니다.",
    "학생이 'teacher' 철자를 틀리지 않으면 기분 좋은 하루.",
    "다음 시험 평균이 오를 기운이 아련하게 느껴집니다.",
    "지금 이 순간에도 누군가는 to와 too를 헷갈리고 있습니다.",
    "오늘의 점심은 평소보다 맛있을 운세입니다.",
    "복도에서 학생이 인사할 확률이 높은 날입니다.",
    "채점하다 보면 은근 손목이 아픕니다. 스트레칭 하세요.",
    "오늘의 행운 숫자는 반 평균과 같습니다.",
    "칠판 글씨가 유난히 잘 써지는 하루가 될 것입니다.",
    "학생이 'I goed'를 안 쓰면 성공적인 하루.",
    "쉬는 시간이 10분이라는 건 선생님한테도 적용됩니다.",
    "오늘의 명언: 가르치는 자가 가장 많이 배운다.",
    "누군가 선생님 책상에 초콜릿을 놓고 갈 기운이 보입니다.",
    "오늘은 출석부 이름을 한 번에 다 맞출 수 있는 날.",
    "프린터가 고장 안 나면 좋은 하루입니다.",
    "다음 시험 범위가 줄어들 기운은 아직 안 보입니다.",
]

# ════════════════════════════════════════════════════════════
# 5. SESSION STATE
# ════════════════════════════════════════════════════════════
for key, default in [
    ("class_data", {}),
    ("excel_bytes", None),
    ("comments_generated", False),
    ("reports_zip", None),
    ("_file_key", None),
    ("_data_ver", 0),
]:
    if key not in st.session_state:
        st.session_state[key] = default


# ════════════════════════════════════════════════════════════
# 6. DOCUMENT UTILITIES (원본 정기테스트_보고서_코드_최종.py 그대로)
# ════════════════════════════════════════════════════════════
def set_table_borders(table, color="E8E8E8", size=3, inside=True, inside_color="E8E8E8", inside_size=3):
    """테이블 테두리 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    for el in tblPr.xpath("./w:tblBorders"):
        tblPr.remove(el)
    borders = OxmlElement('w:tblBorders')

    def make_border(tag, border_color, size_val):
        e = OxmlElement(tag)
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), str(size_val))
        e.set(qn('w:color'), border_color)
        return e

    borders.append(make_border('w:top', color, size))
    borders.append(make_border('w:left', color, size))
    borders.append(make_border('w:bottom', color, size))
    borders.append(make_border('w:right', color, size))

    if inside:
        borders.append(make_border('w:insideH', inside_color, inside_size))
        borders.append(make_border('w:insideV', inside_color, inside_size))

    tblPr.append(borders)

def set_page_border(section, color="4A5A8C", size=24):
    """페이지 전체 테두리 설정"""
    sectPr = section._sectPr

    for border in sectPr.xpath(".//w:pgBorders"):
        sectPr.remove(border)

    pgBorders = OxmlElement('w:pgBorders')
    pgBorders.set(qn('w:offsetFrom'), 'page')

    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), str(size))
        border.set(qn('w:color'), color)
        border.set(qn('w:space'), '0')
        pgBorders.append(border)

    sectPr.append(pgBorders)

def set_cell_background(cell, color="F0F0F0"):
    """셀 배경색 설정"""
    cell_elem = cell._tc
    tcPr = cell_elem.get_or_add_tcPr()

    for shd in tcPr.xpath("./w:shd"):
        tcPr.remove(shd)

    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_cell_padding(table, top=70, left=90, bottom=70, right=90):
    """셀 패딩 설정"""
    tbl = table._tbl
    tblPr = tbl.tblPr

    for el in tblPr.xpath("./w:tblCellMar"):
        tblPr.remove(el)
    cellMar = OxmlElement('w:tblCellMar')

    for side, val in [('top',top), ('left',left), ('bottom',bottom), ('right',right)]:
        elem = OxmlElement('w:'+side)
        elem.set(qn('w:w'), str(val))
        elem.set(qn('w:type'), 'dxa')
        cellMar.append(elem)
    tblPr.append(cellMar)

def _set_east_asia_font(run, font_name='맑은 고딕'):
    """한글(CJK) 폰트를 명시적으로 지정"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)

def style_paragraph(p, size=10, bold=False, align=None, color="1A1A1A"):
    """단락 스타일 설정"""
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = '맑은 고딕'
        _set_east_asia_font(run)
        if color:
            run.font.color.rgb = RGBColor.from_string(color)

    if align:
        p.alignment = align
    p.space_before = Pt(0)
    p.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

def add_label_cell(cell, text, bg_color="F5F5F5"):
    """라벨 셀 추가 (세로 가운데 정렬)"""
    cell.text = text
    set_cell_background(cell, bg_color)
    for p in cell.paragraphs:
        style_paragraph(p, size=10, bold=True, color="1A1A1A")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_value_cell(cell, text="", bg_color="FFFFFF", align_center=True, vertical_center=True):
    """값 셀 추가"""
    if text is None:
        text = ""
    else:
        text = str(text)

    if '\n' in text:
        cell.text = ""
        for line in text.split('\n'):
            p = cell.add_paragraph(line)
            style_paragraph(
                p, size=10, bold=False, color="1A1A1A",
                align=WD_ALIGN_PARAGRAPH.LEFT if not align_center else WD_ALIGN_PARAGRAPH.CENTER
            )
    else:
        cell.text = text
        for p in cell.paragraphs:
            style_paragraph(
                p, size=10, bold=False, color="1A1A1A",
                align=WD_ALIGN_PARAGRAPH.LEFT if not align_center else WD_ALIGN_PARAGRAPH.CENTER
            )

    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER if vertical_center else WD_ALIGN_VERTICAL.TOP

def add_thick_divider(doc, color="4A5A8C", width_cm=19.0, height=60):
    """굵은 구분선 추가"""
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(6)

    divider_table = doc.add_table(rows=1, cols=1)
    divider_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    divider_cell = divider_table.cell(0, 0)
    divider_cell.width = Cm(width_cm)

    divider_cell.text = ""
    set_cell_background(divider_cell, color)

    tbl = divider_table._tbl
    tblPr = tbl.tblPr
    for el in tblPr.xpath("./w:tblBorders"):
        tblPr.remove(el)

    tr = divider_table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(height))
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)

def add_logo_and_title(doc, logo_bytes=None):
    """로고와 제목 추가"""
    title_table = doc.add_table(rows=1, cols=3)
    title_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    tbl = title_table._tbl
    tblPr = tbl.tblPr
    for el in tblPr.xpath("./w:tblBorders"):
        tblPr.remove(el)

    title_table.rows[0].cells[0].width = Cm(4.5)
    title_table.rows[0].cells[1].width = Cm(12.0)
    title_table.rows[0].cells[2].width = Cm(3.0)

    logo_cell = title_table.cell(0, 0)
    logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if logo_bytes:
        try:
            logo_paragraph = logo_cell.paragraphs[0]
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = logo_paragraph.runs[0] if logo_paragraph.runs else logo_paragraph.add_run()
            run.add_picture(io.BytesIO(logo_bytes), width=Cm(4.0))
        except Exception:
            logo_cell.text = ""
    else:
        logo_cell.text = ""

    title_cell = title_table.cell(0, 1)
    title_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run1 = title_paragraph.add_run("최상위 학원\n")
    title_run1.font.size = Pt(24)
    title_run1.font.bold = True
    title_run1.font.name = '맑은 고딕'
    _set_east_asia_font(title_run1)
    title_run1.font.color.rgb = RGBColor.from_string("4A5A8C")

    title_run2 = title_paragraph.add_run("영어 정기 시험 Report")
    title_run2.font.size = Pt(24)
    title_run2.font.bold = True
    title_run2.font.name = '맑은 고딕'
    _set_east_asia_font(title_run2)
    title_run2.font.color.rgb = RGBColor.from_string("4A5A8C")

    right_cell = title_table.cell(0, 2)
    right_cell.text = ""

def remove_trailing_numbers(name):
    """학생 이름 끝에 붙은 숫자 제거"""
    return re.sub(r'\d+$', '', name).strip()

def prevent_table_page_break(table):
    """표가 페이지를 넘어가지 않도록 설정"""
    for row in table.rows:
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cant_split = OxmlElement('w:cantSplit')
        trPr.append(cant_split)


# ════════════════════════════════════════════════════════════
# 7. BUSINESS LOGIC
# ════════════════════════════════════════════════════════════
def _patch_openpyxl():
    """openpyxl의 named_styles 검증을 패치하여 한셀 호환."""
    try:
        from openpyxl.styles.named_styles import _NamedCellStyle
        _orig_init = _NamedCellStyle.__init__
        def _patched_init(self, **kw):
            if kw.get('name') is None:
                kw['name'] = 'Normal'
            _orig_init(self, **kw)
        _NamedCellStyle.__init__ = _patched_init
    except Exception:
        pass

_patch_openpyxl()

def parse_excel(file_bytes):
    class_data = {}
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    sheets = xls.sheet_names
    for cn in ALL_CLASSES:
        if f"{cn}_반정보" not in sheets or f"{cn}_학생" not in sheets:
            continue
        info_df = pd.read_excel(xls, sheet_name=f"{cn}_반정보")
        info = dict(zip(info_df["항목"], info_df["값"]))
        raw = str(info.get("수업시간","") or "")
        if cn == "월수_3부":
            info["수업시간"] = re.sub(r"(월/수)\s*\d{2}:\d{2}-\d{2}:\d{2}", r"\1 20:00-22:30", raw) or "월/수 20:00-22:30"
        elif cn == "화목_3부":
            info["수업시간"] = re.sub(r"(화/목)\s*\d{2}:\d{2}-\d{2}:\d{2}", r"\1 20:00-22:30", raw) or "화/목 20:00-22:30"
        sdf = pd.read_excel(xls, sheet_name=f"{cn}_학생")
        def fc(cands):
            for c in cands:
                if c in sdf.columns: return c
            return None
        cr = fc(["Reading점수","reading점수","리딩점수"])
        cg = fc(["Grammar점수","grammer점수","Grammar점","grammar점","그래머점수"])
        ca = fc(["수업태도"]); cs = fc(["성실성"]); cn_ = fc(["특이사항"])
        ck = fc(["키워드"]); cc = fc(["코멘트","Teacher Comment","teacher comment"])
        students = []
        for idx, row in sdf.iterrows():
            if not (pd.notna(row.get("학생명")) and str(row["학생명"]).strip()):
                continue
            sch = str(row["학교"]) if pd.notna(row.get("학교")) else ""
            gr = str(row["학년"]).replace(".0","") if pd.notna(row.get("학년")) else ""
            sg = f"{sch} {gr}학년" if sch and gr else sch or (f"{gr}학년" if gr else "")
            students.append({
                "학생명": str(row["학생명"]).strip(),
                "학교/학년": sg,
                "Reading점수": int(row[cr]) if cr and pd.notna(row.get(cr)) else 0,
                "Grammar점수": int(row[cg]) if cg and pd.notna(row.get(cg)) else 0,
                "수업태도": str(row[ca]) if ca and pd.notna(row.get(ca)) else "",
                "성실성": str(row[cs]) if cs and pd.notna(row.get(cs)) else "",
                "특이사항": str(row[cn_]) if cn_ and pd.notna(row.get(cn_)) else "",
                "키워드": str(row[ck]) if ck and pd.notna(row.get(ck)) else "",
                "코멘트": str(row[cc]) if cc and pd.notna(row.get(cc)) and str(row[cc]).strip() else "",
                "담당T": info.get("담당T",""),
            })
        class_data[cn] = {"info": info, "students": students}
    return class_data

def calc_avg(students):
    rs = [s["Reading점수"] for s in students if s["Reading점수"] > 0]
    gs = [s["Grammar점수"] for s in students if s["Grammar점수"] > 0]
    return (round(mean(rs)) if rs else 0), (round(mean(gs)) if gs else 0)

def _sanitize(text, max_len=500):
    """프롬프트 인젝션 방어: 제어문자 제거, 길이 제한."""
    if not isinstance(text, str):
        return ""
    text = "".join(c for c in text if ord(c) >= 32 or c in "\n\t")
    return text[:max_len].strip()

def gen_comment(student, r_avg, g_avg, sys_prompt, usr_template):
    try:
        client = openai.OpenAI(api_key=API_KEY)
        kw = _sanitize(student.get("키워드",""), 100)
        ks = f"\n- 핵심 키워드 (반드시 포함): {kw}" if kw else ""
        kg = '8. "키워드"에 적힌 내용을 반드시 코멘트에 자연스럽게 포함시켜야 합니다.' if kw else ""
        prompt = usr_template.format(
            student_name=_sanitize(student.get("학생명",""), 50),
            reading_score=student.get("Reading점수",0),
            grammar_score=student.get("Grammar점수",0), reading_avg=r_avg, grammar_avg=g_avg,
            attitude=_sanitize(student.get("수업태도",""), 100),
            sincerity=_sanitize(student.get("성실성",""), 100),
            notes=_sanitize(student.get("특이사항","")), keyword_section=ks, keyword_guideline=kg,
        )
        resp = client.chat.completions.create(
            model=AI_MODEL,
            messages=[{"role":"system","content":sys_prompt},{"role":"user","content":prompt}],
            max_completion_tokens=AI_MAX_TOKENS, temperature=AI_TEMPERATURE,
        )
        return resp.choices[0].message.content.strip()
    except openai.AuthenticationError:
        return "[오류] API 키가 유효하지 않습니다."
    except openai.RateLimitError:
        return "[오류] API 요청 한도 초과. 잠시 후 다시 시도하세요."
    except Exception:
        return "[오류] 코멘트 생성에 실패했습니다. 네트워크 연결을 확인하세요."

def create_individual_report(student_data, basic_info, reading_avg, grammar_avg, logo_bytes, test_title="2026년도 1차 정기테스트 결과"):
    """개별 학생 보고서 생성 (원본 코드 그대로)"""
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin, section.bottom_margin = Cm(1.0), Cm(1.0)
    section.left_margin, section.right_margin = Cm(1.0), Cm(1.0)

    set_page_border(section)
    doc.add_paragraph()
    add_logo_and_title(doc, logo_bytes)
    add_thick_divider(doc)

    # 학생 정보 테이블
    info_table = doc.add_table(rows=3, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [3.5, 6.0, 3.5, 6.0]
    for i, width in enumerate(widths):
        for row in info_table.rows:
            row.cells[i].width = Cm(width)
    add_label_cell(info_table.cell(0, 0), "학생명")
    display_name = remove_trailing_numbers(student_data['학생명'])
    add_value_cell(info_table.cell(0, 1), display_name, vertical_center=True)
    add_label_cell(info_table.cell(0, 2), "담당T")
    add_value_cell(info_table.cell(0, 3), student_data['담당T'], vertical_center=True)
    add_label_cell(info_table.cell(1, 0), "학교/학년")
    add_value_cell(info_table.cell(1, 1), student_data.get('학교/학년', ''), vertical_center=True)
    add_label_cell(info_table.cell(1, 2), "수업시간")
    add_value_cell(info_table.cell(1, 3), basic_info.get('수업시간', ''), vertical_center=True)
    add_label_cell(info_table.cell(2, 0), "시험일자")
    add_value_cell(info_table.cell(2, 1), str(basic_info.get('시험일자', '')).split(' ')[0], vertical_center=True)
    info_table.cell(2, 2).merge(info_table.cell(2, 3))
    set_table_borders(info_table)
    prevent_table_page_break(info_table)

    add_thick_divider(doc)

    # 시험 결과 테이블
    result_table = doc.add_table(rows=4, cols=3)
    result_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [3.5, 7.75, 7.75]
    for i, width in enumerate(widths):
        for row in result_table.rows:
            row.cells[i].width = Cm(width)
    result_table.cell(0, 0).merge(result_table.cell(0, 2))
    p = result_table.cell(0, 0).paragraphs[0]
    p.add_run(test_title).bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_label_cell(result_table.cell(1, 1), "Reading")
    add_label_cell(result_table.cell(1, 2), "Grammar")
    add_label_cell(result_table.cell(2, 0), "점수")
    add_value_cell(result_table.cell(2, 1), f"{student_data['Reading점수']}/100", vertical_center=True)
    add_value_cell(result_table.cell(2, 2), f"{student_data['Grammar점수']}/100", vertical_center=True)
    add_label_cell(result_table.cell(3, 0), "반평균")
    add_value_cell(result_table.cell(3, 1), f"{reading_avg}/100", vertical_center=True)
    add_value_cell(result_table.cell(3, 2), f"{grammar_avg}/100", vertical_center=True)
    set_table_borders(result_table)
    prevent_table_page_break(result_table)

    # 교재 진도 테이블
    progress_table = doc.add_table(rows=2, cols=4)
    progress_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [3.5, 5.16, 5.16, 5.16]
    for i, width in enumerate(widths):
        for row in progress_table.rows:
            row.cells[i].width = Cm(width)
    progress_table.cell(0, 0).merge(progress_table.cell(1, 0))
    add_label_cell(progress_table.cell(0, 0), "현재\n교재 진도")
    add_label_cell(progress_table.cell(0, 1), "Reading")
    add_label_cell(progress_table.cell(0, 2), "Grammar")
    add_label_cell(progress_table.cell(0, 3), "Listening")
    add_value_cell(progress_table.cell(1, 1), basic_info.get('Reading교재진도', ''), vertical_center=True)
    add_value_cell(progress_table.cell(1, 2), basic_info.get('Grammar교재진도', ''), vertical_center=True)
    add_value_cell(progress_table.cell(1, 3), basic_info.get('Listening교재진도', ''), vertical_center=True)
    set_table_borders(progress_table)
    prevent_table_page_break(progress_table)

    add_thick_divider(doc)

    # 코멘트 테이블
    comment_table = doc.add_table(rows=1, cols=2)
    comment_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    comment_table.rows[0].cells[0].width = Cm(3.5)
    comment_table.rows[0].cells[1].width = Cm(15.5)
    tr = comment_table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(Cm(8.5).twips))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)
    add_label_cell(comment_table.cell(0, 0), "Teacher's\nComment")

    final_comment = student_data.get('코멘트', '')
    if not final_comment:
        final_comment = "코멘트가 입력되지 않았습니다."
    add_value_cell(comment_table.cell(0, 1), final_comment, align_center=False, vertical_center=True)
    set_table_borders(comment_table)
    prevent_table_page_break(comment_table)

    return doc

def doc_to_bytes(doc):
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

def docx_to_pdf_bytes(docx_bytes):
    """Word → PDF. 실패 시 (None, 에러메시지) 반환."""
    with tempfile.TemporaryDirectory() as tmp:
        dp = os.path.join(tmp, "report.docx")
        pp = os.path.join(tmp, "report.pdf")
        with open(dp, "wb") as f:
            f.write(docx_bytes)
        try:
            if platform.system() == "Windows":
                import pythoncom
                pythoncom.CoInitialize()
                try:
                    from docx2pdf import convert
                    convert(dp, pp)
                finally:
                    pythoncom.CoUninitialize()
            else:
                subprocess.run(
                    ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmp, dp],
                    capture_output=True, timeout=60, check=True,
                )
        except Exception:
            return None, "PDF 변환 실패. Word/LibreOffice 설치를 확인하세요."
        if os.path.exists(pp):
            with open(pp, "rb") as f:
                return f.read(), None
    return None, "PDF 파일이 생성되지 않음"

def pdf_to_img_bytes(pdf_bytes, dpi=300):
    """PDF → JPG. 실패 시 (None, 에러메시지) 반환."""
    with tempfile.TemporaryDirectory() as tmp:
        pp = os.path.join(tmp, "report.pdf")
        with open(pp, "wb") as f:
            f.write(pdf_bytes)
        try:
            from pdf2image import convert_from_path
            kw = {"dpi": dpi, "first_page": 1, "last_page": 1}
            if platform.system() == "Windows":
                pop = os.environ.get("POPPLER_PATH", r"C:\Release-25.07.0\poppler-25.07.0\Library\bin")
                if os.path.exists(pop):
                    kw["poppler_path"] = pop
            imgs = convert_from_path(pp, **kw)
            if imgs:
                buf = io.BytesIO()
                imgs[0].save(buf, "JPEG", quality=95)
                return buf.getvalue(), None
            return None, "이미지 변환 결과 없음"
        except Exception:
            return None, "이미지 변환 실패. Poppler 설치를 확인하세요."

def _cc_convert(docx_bytes, output_format, extra_options=None):
    """CloudConvert API로 docx를 지정 형식으로 변환. (bytes, None) 또는 (None, 에러)."""
    if not CC_API_KEY:
        return None, "CloudConvert API 키가 설정되지 않았습니다."
    tmp_path = None
    try:
        import cloudconvert
        import requests as req

        cloudconvert.configure(api_key=CC_API_KEY, sandbox=False)

        convert_opts = {
            'operation': 'convert',
            'input': 'upload',
            'output_format': output_format,
        }
        if extra_options:
            convert_opts.update(extra_options)

        job = cloudconvert.Job.create(payload={
            'tasks': {
                'upload': {'operation': 'import/upload'},
                'convert': convert_opts,
                'export': {
                    'operation': 'export/url',
                    'input': 'convert',
                },
            }
        })

        upload_task_id = None
        export_task_id = None
        for task in job['tasks']:
            if task['operation'] == 'import/upload':
                upload_task_id = task['id']
            elif task['operation'] == 'export/url':
                export_task_id = task['id']

        # upload task를 다시 조회해서 form 정보 획득
        upload_task = cloudconvert.Task.find(id=upload_task_id)

        # 임시파일로 저장 후 업로드
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp.write(docx_bytes)
            tmp_path = tmp.name

        cloudconvert.Task.upload(file_name=tmp_path, task=upload_task)

        # 완료 대기
        res = cloudconvert.Task.wait(id=export_task_id)
        files = res.get('result', {}).get('files', [])
        if not files:
            return None, "변환 결과 없음"

        resp = req.get(files[0]['url'], timeout=60)
        if resp.status_code == 200:
            return resp.content, None
        return None, f"다운로드 실패: {resp.status_code}"
    except Exception as e:
        return None, f"CloudConvert 변환 실패: {e}"
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)

def cloudconvert_docx_to_jpg(docx_bytes):
    return _cc_convert(docx_bytes, 'jpg', {'pixel_density': 300})

def cloudconvert_docx_to_pdf(docx_bytes):
    return _cc_convert(docx_bytes, 'pdf')

def make_template_bytes():
    sample_info = {
        "담당T": "홍길동", "수업시간": "월/수 14:00-15:30", "시험일자": "2025-12-19",
        "Reading교재진도": "YBM 김 Lesson 3-2", "Grammar교재진도": "문법 워크북 Unit 4-1",
        "Listening교재진도": "Listening 교재 Unit 3",
    }
    sample_students = pd.DataFrame([
        {"학생명":"김민준","학교":"대치중","학년":"2","Reading점수":85,"Grammar점수":78,
         "수업태도":"적극적","성실성":"우수","특이사항":"수업 참여도가 높고 질문이 많음","키워드":"독해력","코멘트":""},
        {"학생명":"이서연","학교":"역삼중","학년":"1","Reading점수":72,"Grammar점수":90,
         "수업태도":"보통","성실성":"우수","특이사항":"문법에 강하나 독해 속도 개선 필요","키워드":"문법","코멘트":""},
        {"학생명":"박지호","학교":"대치중","학년":"2","Reading점수":65,"Grammar점수":55,
         "수업태도":"소극적","성실성":"보통","특이사항":"집중력이 떨어지는 편이나 노력하는 모습 보임","키워드":"집중력","코멘트":""},
    ])
    empty_info = {"담당T":"","수업시간":"","시험일자":"","Reading교재진도":"","Grammar교재진도":"","Listening교재진도":""}
    times = {"월수_1부":"월/수 14:00-15:30","월수_2부":"월/수 16:00-17:30","월수_3부":"월/수 20:00-22:30",
             "화목_1부":"화/목 14:00-15:30","화목_2부":"화/목 16:00-17:30","화목_3부":"화/목 20:00-22:30"}
    empty_students = pd.DataFrame({
        "학생명":[""]*20,"학교":[""]*20,"학년":[""]*20,"Reading점수":[""]*20,"Grammar점수":[""]*20,
        "수업태도":[""]*20,"성실성":[""]*20,"특이사항":[""]*20,"키워드":[""]*20,"코멘트":[""]*20,
    })
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        for cn in ALL_CLASSES:
            if cn == "월수_1부":
                pd.DataFrame({"항목":list(sample_info.keys()),"값":list(sample_info.values())}).to_excel(w, sheet_name=f"{cn}_반정보", index=False)
                sample_students.to_excel(w, sheet_name=f"{cn}_학생", index=False)
            else:
                inf = {**empty_info, "수업시간": times[cn]}
                pd.DataFrame({"항목":list(inf.keys()),"값":list(inf.values())}).to_excel(w, sheet_name=f"{cn}_반정보", index=False)
                empty_students.to_excel(w, sheet_name=f"{cn}_학생", index=False)
    return buf.getvalue()

def export_excel(class_data):
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        for cn, data in class_data.items():
            info = data["info"]
            pd.DataFrame({"항목":list(info.keys()),"값":list(info.values())}).to_excel(w, sheet_name=f"{cn}_반정보", index=False)
            if data["students"]:
                cols = ["학생명","학교/학년","Reading점수","Grammar점수","수업태도","성실성","특이사항","키워드","코멘트"]
                df = pd.DataFrame(data["students"])
                df[[c for c in cols if c in df.columns]].to_excel(w, sheet_name=f"{cn}_학생", index=False)
            else:
                pd.DataFrame().to_excel(w, sheet_name=f"{cn}_학생", index=False)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════
# 8. LOGO
# ════════════════════════════════════════════════════════════
def load_default_logo():
    p = Path(__file__).parent / "logo.png"
    return p.read_bytes() if p.exists() else None

DEFAULT_LOGO = load_default_logo()


# ════════════════════════════════════════════════════════════
# 9. SIDEBAR
# ════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("### 보고서 설정")
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    test_title = st.text_input("시험 제목", value="2026년도 1차 정기테스트 결과")

    # 프롬프트
    st.markdown('<p class="section-label">프롬프트 편집</p>', unsafe_allow_html=True)
    with st.expander("시스템 프롬프트"):
        sys_prompt = st.text_area("sys", value=DEFAULT_SYSTEM_PROMPT, height=100, label_visibility="collapsed")
    with st.expander("코멘트 작성 지침"):
        usr_prompt = st.text_area("usr", value=DEFAULT_USER_PROMPT, height=300, label_visibility="collapsed",
            help="변수: {student_name}, {reading_score}, {grammar_score}, {reading_avg}, {grammar_avg}, {attitude}, {sincerity}, {notes}, {keyword_section}, {keyword_guideline}")

    if not API_KEY:
        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.error("API Key 미설정 — .env 파일 또는 Streamlit Secrets에 OPENAI_API_KEY를 설정하세요.")

    # ── 쉬는 시간 ──
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    if st.button("오늘의 강사 운세", use_container_width=True):
        st.info(random.choice(FORTUNES))


# ════════════════════════════════════════════════════════════
# 10. MAIN CONTENT
# ════════════════════════════════════════════════════════════
st.markdown("# 최상위학원 보고서 자동화")
st.caption("영어 정기시험 보고서를 자동으로 생성합니다")
st.markdown('<div class="divider-accent"></div>', unsafe_allow_html=True)

tab1, tab2, tab3 = st.tabs(["데이터 업로드", "코멘트 생성", "보고서 생성"])


# ── Tab 1 ──
with tab1:
    c1, c2 = st.columns([3,1])
    with c2:
        st.download_button(
            "빈 템플릿 다운로드", data=make_template_bytes(),
            file_name="최상위학원_영어시험_6개반_데이터.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )
    with c1:
        uploaded = st.file_uploader("엑셀 파일을 업로드하세요", type=["xlsx","xls"], help="6개 반 데이터가 포함된 엑셀 파일")

    if uploaded:
        file_key = f"{uploaded.name}_{uploaded.size}"
        if st.session_state._file_key != file_key:
            if uploaded.size > 10 * 1024 * 1024:
                st.error("파일 크기가 10MB를 초과합니다.")
            else:
                fb = uploaded.read()
                try:
                    parsed = parse_excel(fb)
                    if not parsed:
                        st.error("유효한 반 데이터를 찾을 수 없습니다. 템플릿 형식을 확인하세요.")
                    else:
                        st.session_state._file_key = file_key
                        st.session_state.excel_bytes = fb
                        st.session_state.class_data = parsed
                        st.session_state.comments_generated = False
                        st.session_state.reports_zip = None
                except Exception as e:
                    st.error(f"엑셀 파일을 읽을 수 없습니다: {e}")

    if st.session_state.class_data:
        cd = st.session_state.class_data
        st.markdown("### 데이터 요약")
        cols = st.columns(min(len(cd), 3))
        for i, (cn, d) in enumerate(cd.items()):
            with cols[i % 3]:
                ns = len(d["students"]); teacher = d["info"].get("담당T","-")
                time_s = d["info"].get("수업시간","-")
                nc = sum(1 for s in d["students"] if s.get("코멘트"))
                tag_cls = "tag-done" if nc == ns else ("tag-partial" if nc > 0 else "tag-empty")
                st.markdown(f"""<div class="card">
                    <div class="card-label">{cn}</div>
                    <div class="card-value">{ns}명</div>
                    <div class="card-meta">{teacher} · {time_s}</div>
                    <div style="margin-top:8px"><span class="tag {tag_cls}">코멘트 {nc}/{ns}</span></div>
                </div>""", unsafe_allow_html=True)

        st.markdown("### 학생 데이터")
        for cn, d in cd.items():
            with st.expander(f"{cn}  ({len(d['students'])}명)"):
                if d["students"]:
                    st.dataframe(
                        pd.DataFrame(d["students"])[["학생명","학교/학년","Reading점수","Grammar점수","수업태도","성실성","특이사항"]],
                        use_container_width=True, hide_index=True,
                    )
                else:
                    st.caption("학생 데이터 없음")
    else:
        st.markdown('<div class="empty">엑셀 파일을 업로드하면 데이터가 여기에 표시됩니다</div>', unsafe_allow_html=True)


# ── Tab 2 ──
with tab2:
    if not st.session_state.class_data:
        st.info("먼저 '데이터 업로드' 탭에서 엑셀 파일을 업로드하세요.")
    elif not API_KEY:
        st.error("API Key가 설정되지 않았습니다. .env 파일을 확인하세요.")
    else:
        cd = st.session_state.class_data
        avail = list(cd.keys())
        selected = st.multiselect("코멘트를 생성할 반", avail, default=avail)

        # 특정 학생 선택
        all_students_t2 = []
        for cn in selected:
            for s in cd[cn]["students"]:
                all_students_t2.append(f"{cn} — {s['학생명']}")
        pick_students_t2 = st.multiselect(
            "특정 학생만 선택 (비워두면 전체)", all_students_t2, default=[], key="pick_t2",
        )

        LOADING_MSGS = [
            "학생 관찰일지 읽는 중...",
            "학부모 감동 포인트 계산 중...",
            "교육적 표현으로 번역 중...",
            "코멘트에 진심을 담는 중...",
            "선생님 필체를 학습하는 중...",
            "부드러운 표현을 고르는 중...",
        ]

        col_btn, col_dl = st.columns([1, 1])
        with col_btn:
            gen_clicked = st.button("코멘트 생성 시작", type="primary", use_container_width=True)
        with col_dl:
            if st.session_state.class_data:
                st.download_button(
                    "코멘트 포함 엑셀 다운로드", data=export_excel(st.session_state.class_data),
                    file_name=f"최상위학원_코멘트완료_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                )

        def _is_picked_t2(cn, name):
            if not pick_students_t2:
                return True
            return f"{cn} — {name}" in pick_students_t2

        if gen_clicked and selected:
            total = sum(1 for cn in selected for s in cd[cn]["students"] if not s.get("코멘트") and _is_picked_t2(cn, s["학생명"]))
            if total == 0:
                st.info("선택된 학생 중 코멘트가 없는 학생이 없습니다.")
            else:
                bar = st.progress(0, text="준비 중...")
                status = st.empty(); fortune_area = st.empty(); done = 0
                for cn in selected:
                    sts = cd[cn]["students"]; ra, ga = calc_avg(sts)
                    for s in sts:
                        if s.get("코멘트"): continue
                        if not _is_picked_t2(cn, s["학생명"]): continue
                        msg = random.choice(LOADING_MSGS)
                        status.markdown(f"**{cn}** · {s['학생명']} — _{msg}_")
                        fortune_area.caption(f"오늘의 운세 — {random.choice(FORTUNES)}")
                        s["코멘트"] = gen_comment(s, ra, ga, sys_prompt, usr_prompt)
                        done += 1; bar.progress(done/total, text=f"{done}/{total}")
                fortune_area.empty()
                bar.progress(1.0, text="완료"); status.success(f"{done}개 코멘트 생성 완료")
                st.balloons()
                st.session_state.comments_generated = True
                st.session_state.reports_zip = None
                st.session_state._data_ver += 1
                st.rerun()

        st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
        st.markdown("### 데이터 수정 / 코멘트 검수")
        st.caption("수업태도, 성실성, 특이사항을 수정한 뒤 코멘트를 생성하세요. 코멘트도 직접 수정 가능합니다.")

        for cn in avail:
            d = cd[cn]
            if not d["students"]: continue
            with st.expander(f"{cn}  ({len(d['students'])}명)", expanded=bool(st.session_state.comments_generated)):
                edf = pd.DataFrame(d["students"])[["학생명","학교/학년","Reading점수","Grammar점수","수업태도","성실성","특이사항","코멘트"]]
                edited = st.data_editor(edf, column_config={
                    "학생명": st.column_config.TextColumn("이름", width=70),
                    "학교/학년": st.column_config.TextColumn("학교/학년", width=90),
                    "Reading점수": st.column_config.NumberColumn("R", width=50),
                    "Grammar점수": st.column_config.NumberColumn("G", width=50),
                    "수업태도": st.column_config.TextColumn("태도", width=60),
                    "성실성": st.column_config.TextColumn("성실성", width=60),
                    "특이사항": st.column_config.TextColumn("특이사항", width=200),
                    "코멘트": st.column_config.TextColumn("코멘트", width=350),
                }, use_container_width=True, hide_index=True, key=f"ed_{cn}_v{st.session_state._data_ver}")
                for idx, row in edited.iterrows():
                    if idx < len(d["students"]):
                        for col in ["학생명","학교/학년","수업태도","성실성","특이사항","코멘트"]:
                            d["students"][idx][col] = row[col] if pd.notna(row[col]) else ""
                        for col in ["Reading점수","Grammar점수"]:
                            d["students"][idx][col] = int(row[col]) if pd.notna(row[col]) else 0



# ── Tab 3 ──
with tab3:
    if not st.session_state.class_data:
        st.info("먼저 '데이터 업로드' 탭에서 엑셀 파일을 업로드하세요.")
    else:
        cd = st.session_state.class_data; avail = list(cd.keys())

        col_cls, col_fmt = st.columns([1,1])
        with col_cls:
            st.markdown("##### 반 선택")
            rpt_classes = st.multiselect("반", avail, default=avail, key="rpt_cls", label_visibility="collapsed")
        with col_fmt:
            st.markdown("##### 출력 형식")
            fc1, fc2, fc3 = st.columns(3)
            with fc1: out_docx = st.checkbox("Word", value=True)
            with fc2: out_pdf  = st.checkbox("PDF", value=False)
            with fc3: out_jpg  = st.checkbox("이미지", value=True)

        # 특정 학생 선택
        all_students_t3 = []
        for cn in rpt_classes:
            for s in cd[cn]["students"]:
                if s.get("코멘트"):
                    all_students_t3.append(f"{cn} — {s['학생명']}")
        pick_students_t3 = st.multiselect(
            "특정 학생만 선택 (비워두면 전체)", all_students_t3, default=[], key="pick_t3",
        )

        def _is_picked_t3(cn, name):
            if not pick_students_t3:
                return True
            return f"{cn} — {name}" in pick_students_t3

        # 준비 상태 확인
        ready = 0; missing = []
        for cn in rpt_classes:
            for s in cd[cn]["students"]:
                if s.get("코멘트") and _is_picked_t3(cn, s["학생명"]): ready += 1
                elif not s.get("코멘트"): missing.append(f"{cn} — {s['학생명']}")
        if missing:
            with st.expander(f"코멘트 없는 학생 {len(missing)}명 (건너뜀)"):
                for m in missing: st.text(m)

        if ready > 0 and (out_docx or out_pdf or out_jpg):
            st.caption(f"총 {ready}명의 보고서를 생성합니다")

            if st.button("보고서 생성", type="primary", use_container_width=True):
                bar = st.progress(0, text="준비 중..."); status = st.empty(); fortune_area2 = st.empty()
                zbuf = io.BytesIO(); done = 0; img_fail = 0; errors = []

                try:
                    with zipfile.ZipFile(zbuf, "w", zipfile.ZIP_DEFLATED) as zf:
                        for cn in rpt_classes:
                            d = cd[cn]; info = d["info"]; sts = d["students"]
                            ra, ga = calc_avg(sts)
                            teacher = info.get("담당T","선생님")
                            ds = datetime.now().strftime("%Y_%m%d")
                            prefix = f"{teacher}_정기테스트_보고서_{ds}"

                            for s in sts:
                                if not s.get("코멘트"): continue
                                if not _is_picked_t3(cn, s["학생명"]): continue
                                try:
                                    status.markdown(f"**{cn}** · {s['학생명']}")
                                    fortune_area2.caption(f"오늘의 운세 — {random.choice(FORTUNES)}")
                                    doc = create_individual_report(s, info, ra, ga, DEFAULT_LOGO, test_title)
                                    dbytes = doc_to_bytes(doc)
                                    safe = re.sub(r'[\\/*?:"<>|]', "_", s["학생명"])
                                    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                                    base = f"{safe}_{ts}"

                                    if out_docx:
                                        zf.writestr(f"{prefix}/{cn}/{base}.docx", dbytes)

                                    if out_pdf:
                                        pbytes, pdf_err = cloudconvert_docx_to_pdf(dbytes)
                                        if pbytes:
                                            zf.writestr(f"{prefix}/{cn}/{base}.pdf", pbytes)
                                        else:
                                            img_fail += 1
                                            errors.append(f"{s['학생명']}: {pdf_err}")

                                    if out_jpg:
                                        ibytes, img_err = cloudconvert_docx_to_jpg(dbytes)
                                        if ibytes:
                                            zf.writestr(f"{prefix}/{cn}/{base}.jpg", ibytes)
                                        else:
                                            img_fail += 1
                                            errors.append(f"{s['학생명']}: {img_err}")

                                    done += 1; bar.progress(done/ready, text=f"{done}/{ready}")
                                    import time; time.sleep(1)
                                except Exception:
                                    errors.append(f"{s['학생명']}: 보고서 생성 실패")

                    fortune_area2.empty()
                    bar.progress(1.0, text="완료")
                    if errors:
                        st.error("일부 보고서 생성 실패:\n" + "\n".join(errors))
                    if img_fail > 0:
                        status.warning(f"보고서 {done}개 생성 완료 — PDF/이미지 변환 실패 {img_fail}건 (Word 파일은 정상 포함)")
                    elif done > 0:
                        status.success(f"보고서 {done}개 생성 완료")
                        st.snow()
                    else:
                        status.warning("코멘트가 있는 학생이 없습니다. 먼저 코멘트를 생성하세요.")
                    st.session_state.reports_zip = zbuf.getvalue()
                except Exception:
                    st.error("보고서 생성 중 오류가 발생했습니다.")

        if st.session_state.reports_zip:
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            teacher = ""
            for cn in avail:
                t = cd[cn]["info"].get("담당T","")
                if t: teacher = t; break
            ds = datetime.now().strftime("%Y_%m%d")
            fname = f"{teacher}_정기테스트_보고서_{ds}.zip" if teacher else f"보고서_{ds}.zip"
            st.download_button(
                "전체 보고서 다운로드 (ZIP)", data=st.session_state.reports_zip,
                file_name=fname, mime="application/zip",
                use_container_width=True, type="primary",
            )
